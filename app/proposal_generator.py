# backend/app/proposal_generator.py
"""
Professional Proposal Generator for SentriBiD.
Creates beautiful, branded PDF and DOCX proposals.
"""

import os
import json
import uuid
from datetime import datetime

UPLOAD_DIR = os.getenv("SENTRIBID_UPLOAD_DIR", "./uploads/opportunities")


def tryparse(s):
    if not s:
        return None
    if isinstance(s, (dict, list)):
        return s
    try:
        return json.loads(s)
    except Exception:
        return None


def generate_pdf_proposal(opp, proposal_text: str, company_name: str, profile=None, bid=None):
    """Generate a professionally formatted PDF proposal."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, white, black
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
        HRFlowable, KeepTogether
    )

    output_dir = os.path.join(UPLOAD_DIR, "proposals")
    os.makedirs(output_dir, exist_ok=True)
    filename = f"proposal_{opp.opp_code}_{uuid.uuid4().hex[:6]}.pdf"
    filepath = os.path.join(output_dir, filename)

    # Brand colors
    NAVY = HexColor('#1B2A4A')
    GOLD = HexColor('#C8A951')
    DARK_GRAY = HexColor('#2C2C2C')
    MED_GRAY = HexColor('#666666')
    LIGHT_GRAY = HexColor('#F5F5F5')
    ACCENT_BLUE = HexColor('#2E6DA4')

    doc = SimpleDocTemplate(
        filepath, pagesize=letter,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='CoverTitle', fontSize=30, leading=36, spaceAfter=4,
        textColor=NAVY, fontName='Helvetica-Bold', alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name='CoverSubtitle', fontSize=14, leading=18, spaceAfter=12,
        textColor=GOLD, fontName='Helvetica-Bold', alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name='CoverInfo', fontSize=11, leading=14, spaceAfter=4,
        textColor=MED_GRAY, fontName='Helvetica', alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name='SectionHead', fontSize=16, leading=20, spaceBefore=20, spaceAfter=10,
        textColor=NAVY, fontName='Helvetica-Bold', borderPadding=(0, 0, 4, 0),
    ))
    styles.add(ParagraphStyle(
        name='SubHead', fontSize=12, leading=15, spaceBefore=12, spaceAfter=6,
        textColor=ACCENT_BLUE, fontName='Helvetica-Bold',
    ))
    styles.add(ParagraphStyle(
        name='ProposalBody', fontSize=10.5, leading=15, spaceAfter=8,
        textColor=DARK_GRAY, fontName='Helvetica', alignment=TA_JUSTIFY,
    ))
    styles.add(ParagraphStyle(
        name='TableHeader', fontSize=10, leading=13, spaceAfter=0,
        textColor=white, fontName='Helvetica-Bold', alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name='TableCell', fontSize=10, leading=13, spaceAfter=0,
        textColor=DARK_GRAY, fontName='Helvetica',
    ))
    styles.add(ParagraphStyle(
        name='TableCellRight', fontSize=10, leading=13, spaceAfter=0,
        textColor=DARK_GRAY, fontName='Helvetica', alignment=TA_RIGHT,
    ))
    styles.add(ParagraphStyle(
        name='Footer', fontSize=8, leading=10, textColor=MED_GRAY,
        fontName='Helvetica', alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name='Confidential', fontSize=8, leading=10, textColor=GOLD,
        fontName='Helvetica-Bold', alignment=TA_CENTER,
    ))

    story = []

    # ─── COVER PAGE ───────────────────────────────────────
    story.append(Spacer(1, 1.2 * inch))

    # Gold line
    story.append(HRFlowable(width="60%", thickness=2, color=GOLD, spaceAfter=20))

    story.append(Paragraph(company_name.upper(), styles['CoverTitle']))
    story.append(Spacer(1, 6))
    story.append(Paragraph("PROPOSAL RESPONSE", styles['CoverSubtitle']))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="40%", thickness=1, color=GOLD, spaceAfter=20))

    story.append(Paragraph(f"<b>{opp.title}</b>", styles['CoverInfo']))
    story.append(Spacer(1, 20))

    # Metadata box
    meta_data = [
        ["Agency:", opp.agency_name or "—"],
        ["Solicitation #:", opp.solicitation_number or "—"],
        ["NAICS Code:", opp.naics_code or "—"],
        ["Due Date:", str(opp.due_date or "—")],
        ["Submitted By:", company_name],
        ["Date:", datetime.now().strftime("%B %d, %Y")],
    ]
    meta_table = Table(meta_data, colWidths=[1.8 * inch, 4.0 * inch])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), NAVY),
        ('TEXTCOLOR', (1, 0), (1, -1), DARK_GRAY),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('LINEBELOW', (0, 0), (-1, -2), 0.5, HexColor('#DDDDDD')),
    ]))
    story.append(meta_table)

    story.append(Spacer(1, 40))
    story.append(Paragraph("CONFIDENTIAL — FOR AUTHORIZED USE ONLY", styles['Confidential']))

    story.append(PageBreak())

    # ─── TABLE OF CONTENTS ────────────────────────────────
    story.append(Paragraph("TABLE OF CONTENTS", styles['SectionHead']))
    story.append(HRFlowable(width="100%", thickness=1.5, color=NAVY, spaceAfter=12))

    toc_items = [
        "1. Cover Letter",
        "2. Executive Summary",
        "3. Technical Approach",
        "4. Management Plan",
        "5. Past Performance",
        "6. Staffing Plan",
        "7. Quality Assurance",
        "8. Pricing",
    ]
    for item in toc_items:
        story.append(Paragraph(item, ParagraphStyle(
            'TOCItem', parent=styles['ProposalBody'],
            fontSize=11, leading=20, leftIndent=20, textColor=NAVY,
        )))
    story.append(PageBreak())

    # ─── PROPOSAL BODY ────────────────────────────────────
    sections = proposal_text.split("\n")
    section_number = 0

    for line in sections:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        if line == "---":
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#DDDDDD'), spaceBefore=8, spaceAfter=8))
            continue

        # Detect headers
        clean = line.replace("**", "").replace("##", "").replace("#", "").strip()
        is_header = False
        is_subheader = False

        if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")):
            is_header = True
            clean = line[2:].strip().replace("**", "").replace(":", "").strip()
            section_number += 1
        elif line.startswith("**") and line.endswith("**"):
            is_subheader = True
            clean = line.replace("**", "").strip()
        elif line.startswith("## "):
            is_header = True
            clean = line.replace("##", "").strip()
        elif clean.upper() == clean and len(clean) > 3 and clean[0].isalpha() and not any(c in clean for c in '.,$'):
            is_header = True

        if is_header:
            story.append(Spacer(1, 8))
            story.append(Paragraph(clean.upper(), styles['SectionHead']))
            story.append(HRFlowable(width="100%", thickness=1.5, color=NAVY, spaceAfter=8))
        elif is_subheader:
            story.append(Paragraph(clean, styles['SubHead']))
        else:
            # Clean markdown formatting for PDF
            text = line.replace("**", "").replace("__", "").replace("*", "")
            try:
                story.append(Paragraph(text, styles['ProposalBody']))
            except Exception:
                story.append(Paragraph(line.encode('ascii', 'replace').decode(), styles['ProposalBody']))

    # ─── PRICING TABLE ────────────────────────────────────
    if bid or opp.ai_bid_strategy:
        story.append(PageBreak())
        story.append(Paragraph("PRICING", styles['SectionHead']))
        story.append(HRFlowable(width="100%", thickness=1.5, color=NAVY, spaceAfter=12))

        # Try to build pricing from bid or AI analysis
        rec = tryparse(opp.ai_bid_strategy) if not bid else None

        if bid:
            # Get pricing from actual bid
            pricing_rows = [
                [Paragraph("Cost Category", styles['TableHeader']),
                 Paragraph("Amount", styles['TableHeader'])],
            ]

            # Item costs
            item_total = sum(getattr(i, 'quantity', 0) * getattr(i, 'unit_cost', 0) for i in (bid.items or []))
            if item_total > 0:
                pricing_rows.append([
                    Paragraph("Materials / Items", styles['TableCell']),
                    Paragraph(f"${item_total:,.2f}", styles['TableCellRight']),
                ])

            # Labor costs
            labor_total = sum(
                getattr(l, 'hourly_rate', 0) * getattr(l, 'hours', 0) * getattr(l, 'workers', 1)
                for l in (bid.labor_lines or [])
            )
            if labor_total > 0:
                pricing_rows.append([
                    Paragraph("Labor", styles['TableCell']),
                    Paragraph(f"${labor_total:,.2f}", styles['TableCellRight']),
                ])

            # Overhead
            if bid.overhead:
                oh = bid.overhead
                oh_total = (
                    getattr(oh, 'insurance_allocation', 0) +
                    getattr(oh, 'admin_time_cost', 0) +
                    getattr(oh, 'bonding_compliance_cost', 0) +
                    getattr(oh, 'misc_overhead', 0)
                )
                if oh_total > 0:
                    pricing_rows.append([
                        Paragraph("Overhead & Administration", styles['TableCell']),
                        Paragraph(f"${oh_total:,.2f}", styles['TableCellRight']),
                    ])

            total = item_total + labor_total
            pricing_rows.append([
                Paragraph("<b>TOTAL ESTIMATED COST</b>", styles['TableCell']),
                Paragraph(f"<b>${total:,.2f}</b>", styles['TableCellRight']),
            ])

            ptable = Table(pricing_rows, colWidths=[4.5 * inch, 2.5 * inch])
            ptable.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), NAVY),
                ('TEXTCOLOR', (0, 0), (-1, 0), white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                ('LINEBELOW', (0, 0), (-1, -2), 0.5, HexColor('#DDDDDD')),
                ('LINEBELOW', (0, -1), (-1, -1), 1.5, NAVY),
                ('BACKGROUND', (0, -1), (-1, -1), HexColor('#F0F4F8')),
            ]))
            story.append(ptable)
        else:
            # Use AI recommendation pricing
            rec_data = tryparse(opp.ai_bid_strategy)
            if opp.ai_confidence_score:
                story.append(Paragraph(
                    f"Based on our analysis, the estimated contract value range and our recommended pricing approach:",
                    styles['ProposalBody']
                ))
                bid_rec = tryparse(opp.ai_bid_recommendation) if isinstance(opp.ai_bid_recommendation, str) else None

        story.append(Spacer(1, 12))
        strategy = tryparse(opp.ai_bid_strategy)
        if strategy and strategy.get("pricing_approach"):
            story.append(Paragraph(
                f"<b>Pricing Approach:</b> {strategy['pricing_approach'].title()}",
                styles['ProposalBody']
            ))
            if strategy.get("pricing_reasoning"):
                story.append(Paragraph(strategy["pricing_reasoning"], styles['ProposalBody']))

    # ─── CLOSING ──────────────────────────────────────────
    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="40%", thickness=1, color=GOLD, spaceAfter=12))
    story.append(Paragraph(
        f"{company_name} — Confidential Proposal | {datetime.now().strftime('%B %Y')}",
        styles['Footer']
    ))

    # Build with page numbers
    def add_page_number(canvas, doc):
        canvas.saveState()
        # Header line
        canvas.setStrokeColor(NAVY)
        canvas.setLineWidth(0.5)
        canvas.line(54, letter[1] - 36, letter[0] - 54, letter[1] - 36)
        # Header text
        canvas.setFont('Helvetica', 7)
        canvas.setFillColor(MED_GRAY)
        canvas.drawString(54, letter[1] - 32, company_name.upper())
        canvas.drawRightString(letter[0] - 54, letter[1] - 32, "CONFIDENTIAL")
        # Footer
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(MED_GRAY)
        canvas.drawCentredString(letter[0] / 2, 30, f"Page {doc.page}")
        # Footer line
        canvas.line(54, 42, letter[0] - 54, 42)
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return filepath


def generate_docx_proposal(opp, proposal_text: str, company_name: str, profile=None, bid=None):
    """Generate a professionally formatted DOCX proposal."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    output_dir = os.path.join(UPLOAD_DIR, "proposals")
    os.makedirs(output_dir, exist_ok=True)
    filename = f"proposal_{opp.opp_code}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(output_dir, filename)

    NAVY = RGBColor(27, 42, 74)
    GOLD = RGBColor(200, 169, 81)
    DARK = RGBColor(44, 44, 44)
    GRAY = RGBColor(102, 102, 102)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK

    # Heading styles
    for level, size, color in [(1, 18, NAVY), (2, 14, NAVY), (3, 12, RGBColor(46, 109, 164))]:
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True

    # ─── COVER PAGE ───────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name.upper())
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = 'Calibri'

    # Gold separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD

    # Proposal Response
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROPOSAL RESPONSE")
    run.font.size = Pt(16)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Opportunity title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(opp.title)
    run.font.size = Pt(16)
    run.font.color.rgb = DARK
    run.bold = True

    doc.add_paragraph()

    # Metadata table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Agency:", opp.agency_name or "—"),
        ("Solicitation #:", opp.solicitation_number or "—"),
        ("NAICS Code:", opp.naics_code or "—"),
        ("Due Date:", str(opp.due_date or "—")),
        ("Submitted By:", company_name),
        ("Date:", datetime.now().strftime("%B %d, %Y")),
    ]
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        # Key cell
        cell_k = row.cells[0]
        cell_k.text = ""
        p = cell_k.paragraphs[0]
        run = p.add_run(k)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run.font.name = 'Calibri'
        # Value cell
        cell_v = row.cells[1]
        cell_v.text = ""
        p = cell_v.paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'

    doc.add_paragraph()

    # Confidential notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — FOR AUTHORIZED USE ONLY")
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD
    run.bold = True

    doc.add_page_break()

    # ─── TABLE OF CONTENTS ────────────────────────────────
    doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_items = [
        "1. Cover Letter", "2. Executive Summary", "3. Technical Approach",
        "4. Management Plan", "5. Past Performance", "6. Staffing Plan",
        "7. Quality Assurance", "8. Pricing",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_page_break()

    # ─── PROPOSAL BODY ────────────────────────────────────
    lines = proposal_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line or line == "---":
            continue

        clean = line.replace("**", "").replace("##", "").replace("#", "").strip()
        is_header = False

        if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")):
            is_header = True
            clean = line[2:].strip().replace("**", "").replace(":", "").strip()
        elif line.startswith("**") and line.endswith("**"):
            is_header = True
            clean = line.replace("**", "").strip()
        elif line.startswith("## "):
            is_header = True
            clean = line.replace("##", "").strip()
        elif clean.upper() == clean and len(clean) > 3 and clean[0].isalpha() and not any(c in clean for c in '.,$'):
            is_header = True

        if is_header:
            doc.add_heading(clean, level=2)
        else:
            p = doc.add_paragraph(line.replace("**", "").replace("__", ""))
            p.paragraph_format.space_after = Pt(6)

    # ─── PRICING TABLE ────────────────────────────────────
    if bid or opp.ai_bid_strategy:
        doc.add_page_break()
        doc.add_heading("PRICING", level=1)

        if bid:
            item_total = sum(
                getattr(i, 'quantity', 0) * getattr(i, 'unit_cost', 0) for i in (bid.items or [])
            )
            labor_total = sum(
                getattr(l, 'hourly_rate', 0) * getattr(l, 'hours', 0) * getattr(l, 'workers', 1)
                for l in (bid.labor_lines or [])
            )
            oh_total = 0
            if bid.overhead:
                oh = bid.overhead
                oh_total = sum([
                    getattr(oh, 'insurance_allocation', 0),
                    getattr(oh, 'admin_time_cost', 0),
                    getattr(oh, 'bonding_compliance_cost', 0),
                    getattr(oh, 'misc_overhead', 0),
                ])

            rows = [("Cost Category", "Amount")]
            if item_total > 0:
                rows.append(("Materials / Items", f"${item_total:,.2f}"))
            if labor_total > 0:
                rows.append(("Labor", f"${labor_total:,.2f}"))
            if oh_total > 0:
                rows.append(("Overhead & Administration", f"${oh_total:,.2f}"))
            total = item_total + labor_total + oh_total
            rows.append(("TOTAL ESTIMATED COST", f"${total:,.2f}"))

            table = doc.add_table(rows=len(rows), cols=2)
            table.style = 'Light Grid Accent 1'
            for i, (cat, amt) in enumerate(rows):
                table.rows[i].cells[0].text = cat
                table.rows[i].cells[1].text = amt
                if i == 0 or i == len(rows) - 1:
                    for cell in table.rows[i].cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True

        strategy = tryparse(opp.ai_bid_strategy)
        if strategy:
            doc.add_paragraph()
            if strategy.get("pricing_approach"):
                p = doc.add_paragraph()
                run = p.add_run("Pricing Approach: ")
                run.bold = True
                p.add_run(strategy["pricing_approach"].title())
            if strategy.get("pricing_reasoning"):
                doc.add_paragraph(strategy["pricing_reasoning"])

    # ─── CLOSING ──────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 20)
    run.font.color.rgb = GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name} — Confidential Proposal | {datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.save(filepath)
    return filepath
