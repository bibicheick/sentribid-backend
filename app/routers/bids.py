# backend/app/routers/bids.py
import json
import uuid
import io
import csv
import os
import shutil
from datetime import datetime, timezone, date
from typing import Optional, List
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import func

from ..db import get_db
from ..auth import require_auth
from .. import models, schemas
from ..copilot import analyze_bid_risk, analyze_bid_profit, analyze_bid_compliance, chat_with_copilot, portfolio_insights, is_ai_enabled

# Attachment storage directory
UPLOAD_DIR = Path(os.getenv("SENTRIBID_UPLOAD_DIR", "./uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "png", "jpg", "jpeg", "gif", "txt", "rtf"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

router = APIRouter(prefix="/bids", tags=["bids"])


# ─── Helpers ───────────────────────────────────────────────

def _generate_bid_code(db: Session) -> str:
    """
    Generate a unique bid code using UUID suffix.
    Format: SB-YYYY-XXXXXXXX (8-char hex from uuid4).
    Checks DB to avoid collisions with old sequential codes.
    """
    year = datetime.now(timezone.utc).year
    for _ in range(20):
        suffix = uuid.uuid4().hex[:8].upper()
        code = f"SB-{year}-{suffix}"
        exists = db.query(models.Bid.id).filter(models.Bid.bid_code == code).first()
        if not exists:
            return code
    # Absolute fallback — 16-char hex, mathematically impossible to collide
    return f"SB-{year}-{uuid.uuid4().hex[:16].upper()}"


def _safe_float(val, default=0.0):
    try:
        return float(val) if val is not None else default
    except (ValueError, TypeError):
        return default


def _compute_totals(bid: models.Bid) -> dict:
    """Compute all cost totals for a bid."""
    item_subtotal = sum(_safe_float(it.quantity) * _safe_float(it.unit_cost) for it in bid.items)

    labor_total = sum(
        _safe_float(lb.hourly_rate) * _safe_float(lb.hours) * _safe_float(lb.workers, 1)
        for lb in bid.labor_lines
    )

    transport_total = 0.0
    if bid.transport:
        t = bid.transport
        per_trip = (
            _safe_float(t.truck_rental_cost)
            + _safe_float(t.fuel_cost)
            + _safe_float(t.mileage_cost)
            + _safe_float(t.toll_fees)
            + _safe_float(t.driver_cost)
        )
        transport_total = per_trip * _safe_float(t.trips, 1)

    equipment_total = sum(
        _safe_float(eq.rental_cost) * _safe_float(eq.rental_days, 1)
        + (_safe_float(eq.operator_cost) if eq.operator_required else 0.0)
        for eq in bid.equipment_lines
    )

    overhead_total = 0.0
    if bid.overhead:
        o = bid.overhead
        overhead_total = (
            _safe_float(o.insurance_allocation)
            + _safe_float(o.storage_cost)
            + _safe_float(o.admin_time_cost)
            + _safe_float(o.bonding_compliance_cost)
            + _safe_float(o.misc_overhead)
        )

    true_cost = item_subtotal + labor_total + transport_total + equipment_total + overhead_total

    # Risk buffer: risk_level * 2% of true_cost
    risk_pct = _safe_float(bid.risk_level, 1) * 0.02
    risk_buffer = true_cost * risk_pct
    adjusted_cost = true_cost + risk_buffer

    return {
        "item_subtotal": round(item_subtotal, 2),
        "labor_total": round(labor_total, 2),
        "transport_total": round(transport_total, 2),
        "equipment_total": round(equipment_total, 2),
        "overhead_total": round(overhead_total, 2),
        "true_cost": round(true_cost, 2),
        "risk_buffer": round(risk_buffer, 2),
        "adjusted_cost": round(adjusted_cost, 2),
    }


def _base_margin(bid: models.Bid) -> float:
    """Determine base margin % from competition and urgency."""
    comp = (bid.competition_level or "medium").lower()
    base = {"low": 18.0, "medium": 14.0, "high": 10.0}.get(comp, 14.0)

    # Override if set
    if bid.margin_override_pct is not None and bid.margin_override_pct > 0:
        base = float(bid.margin_override_pct)

    # Urgency adjustment
    urgency = _safe_float(bid.urgency_level, 3)
    if urgency >= 4:
        base += 2.0
    elif urgency <= 2:
        base -= 1.0

    return round(max(base, 3.0), 2)


def _recommendations(totals: dict, base_margin: float, bid: models.Bid) -> list:
    """Generate 3 pricing recommendations: conservative, balanced, aggressive."""
    adjusted = totals["adjusted_cost"]
    min_profit = _safe_float(bid.min_acceptable_profit, 0)

    modes = {
        "conservative": base_margin - 3.0,
        "balanced": base_margin,
        "aggressive": base_margin + 4.0,
    }

    recs = []
    for mode, margin in modes.items():
        margin = max(margin, 3.0)
        bid_price = adjusted * (1 + margin / 100)
        profit = bid_price - adjusted
        warnings = []

        if profit < min_profit and min_profit > 0:
            warnings.append(f"Profit ${profit:.0f} below minimum ${min_profit:.0f}")

        # Win score heuristic (0-100)
        comp_factor = {"low": 85, "medium": 70, "high": 55}.get(
            (bid.competition_level or "medium").lower(), 70
        )
        margin_penalty = max(0, (margin - 15) * 2)
        win_score = max(10, min(100, int(comp_factor - margin_penalty + (_safe_float(bid.urgency_level, 3) * 2))))

        recs.append({
            "mode": mode,
            "margin_pct": round(margin, 2),
            "bid_price": round(bid_price, 2),
            "profit_amount": round(profit, 2),
            "win_score": win_score,
            "warnings": warnings,
        })

    return recs


# ─── CRUD ──────────────────────────────────────────────────

@router.get("", response_model=List[schemas.BidOut])
def list_bids(
    q: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    query = db.query(models.Bid)
    if status:
        query = query.filter(models.Bid.status == status.lower())
    if q:
        like = f"%{q}%"
        query = query.filter(
            models.Bid.contract_title.ilike(like)
            | models.Bid.agency_name.ilike(like)
            | models.Bid.bid_code.ilike(like)
        )
    query = query.order_by(models.Bid.id.desc())
    return query.offset(offset).limit(limit).all()


@router.post("", response_model=schemas.BidOut, status_code=201)
def create_bid(
    payload: schemas.BidCreate,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    from sqlalchemy.exc import IntegrityError

    def _make_bid(code: str) -> models.Bid:
        return models.Bid(
            bid_code=code,
            contract_title=payload.contract_title,
            agency_name=payload.agency_name,
            agency_type=payload.agency_type.lower(),
            solicitation_number=payload.solicitation_number or None,
            procurement_method=(payload.procurement_method or "").lower() or None,
            contract_type=payload.contract_type.lower(),
            delivery_distance_miles=payload.delivery_distance_miles,
            deadline_date=payload.deadline_date,
            urgency_level=payload.urgency_level,
            competition_level=payload.competition_level.lower(),
            risk_level=payload.risk_level,
            desired_profit_mode=payload.desired_profit_mode.lower(),
            min_acceptable_profit=payload.min_acceptable_profit,
            margin_override_pct=payload.margin_override_pct,
            notes=payload.notes or None,
            status="draft",
        )

    # Try up to 3 times (collision is virtually impossible with 8-char hex)
    for attempt in range(3):
        bid_code = _generate_bid_code(db)
        bid = _make_bid(bid_code)
        try:
            db.add(bid)
            db.commit()
            db.refresh(bid)
            return bid
        except IntegrityError:
            db.rollback()
            if attempt == 2:
                raise HTTPException(
                    status_code=409,
                    detail=f"Could not generate a unique bid code after 3 attempts. Please try again."
                )
            # Loop will retry with a new code

    raise HTTPException(status_code=500, detail="Unexpected error creating bid.")


# ─── Static path routes MUST come before /{bid_id} ────────

@router.get("/copilot/status")
def copilot_status(user: str = Depends(require_auth)):
    """Check if OpenAI is configured."""
    return {"ai_enabled": is_ai_enabled(), "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini")}

@router.get("/copilot/portfolio")
def copilot_portfolio(
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Get AI portfolio-level insights across all bids."""
    bids = db.query(models.Bid).all()
    bids_data = []
    for b in bids:
        bids_data.append({
            "id": b.id,
            "bid_code": b.bid_code,
            "contract_title": b.contract_title,
            "agency_name": b.agency_name,
            "agency_type": b.agency_type,
            "contract_type": b.contract_type,
            "competition_level": b.competition_level,
            "risk_level": b.risk_level,
            "urgency_level": b.urgency_level,
            "status": b.status,
            "deadline_date": str(b.deadline_date) if b.deadline_date else None,
        })
    return portfolio_insights(bids_data)


# ─── Dynamic {bid_id} routes below ────────────────────────

@router.get("/{bid_id}", response_model=schemas.BidOut)
def get_bid(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")
    return bid


@router.delete("/{bid_id}")
def delete_bid(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Delete a bid and all related data (items, labor, transport, overhead, equipment, versions, outcome)."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")
    if bid.status == "approved":
        raise HTTPException(status_code=400, detail="Cannot delete an approved bid. Revert to draft first.")
    db.delete(bid)
    db.commit()
    return {"ok": True, "deleted_id": bid_id}


@router.patch("/{bid_id}", response_model=schemas.BidOut)
def update_bid(
    bid_id: int,
    payload: dict,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Update editable fields on a bid. Accepts partial updates."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    allowed = {
        "contract_title", "agency_name", "agency_type", "solicitation_number",
        "procurement_method", "contract_type", "delivery_distance_miles",
        "deadline_date", "urgency_level", "competition_level", "risk_level",
        "desired_profit_mode", "min_acceptable_profit", "margin_override_pct", "notes",
    }
    lower_fields = {"agency_type", "procurement_method", "contract_type", "competition_level", "desired_profit_mode"}

    for key, val in payload.items():
        if key not in allowed:
            continue
        if key in lower_fields and isinstance(val, str):
            val = val.lower()
        if key == "deadline_date" and isinstance(val, str):
            from datetime import date as date_type
            try:
                val = date_type.fromisoformat(val[:10])
            except (ValueError, TypeError):
                continue
        setattr(bid, key, val)

    db.commit()
    db.refresh(bid)
    return bid


@router.patch("/{bid_id}/items/{item_id}")
def update_item(
    bid_id: int,
    item_id: int,
    payload: dict,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Update fields on a line item."""
    item = db.query(models.BidItem).filter(
        models.BidItem.id == item_id, models.BidItem.bid_id == bid_id
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")

    allowed = {"name", "description", "quantity", "unit_cost", "supplier_name", "risk_flag"}
    for key, val in payload.items():
        if key in allowed:
            setattr(item, key, val)

    db.commit()
    db.refresh(item)
    return {"ok": True, "id": item.id}


@router.patch("/{bid_id}/labor/{labor_id}")
def update_labor(
    bid_id: int,
    labor_id: int,
    payload: dict,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Update fields on a labor line."""
    labor = db.query(models.BidLaborLine).filter(
        models.BidLaborLine.id == labor_id, models.BidLaborLine.bid_id == bid_id
    ).first()
    if not labor:
        raise HTTPException(status_code=404, detail="Labor line not found")

    allowed = {"labor_type", "hourly_rate", "hours", "workers"}
    for key, val in payload.items():
        if key in allowed:
            setattr(labor, key, val)

    db.commit()
    db.refresh(labor)
    return {"ok": True, "id": labor.id}


@router.delete("/{bid_id}/labor/{labor_id}")
def delete_labor(
    bid_id: int,
    labor_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    labor = db.query(models.BidLaborLine).filter(
        models.BidLaborLine.id == labor_id, models.BidLaborLine.bid_id == bid_id
    ).first()
    if not labor:
        raise HTTPException(status_code=404, detail="Labor line not found")
    db.delete(labor)
    db.commit()
    return {"ok": True}


@router.delete("/{bid_id}/equipment/{equip_id}")
def delete_equipment(
    bid_id: int,
    equip_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    eq = db.query(models.BidEquipmentLine).filter(
        models.BidEquipmentLine.id == equip_id, models.BidEquipmentLine.bid_id == bid_id
    ).first()
    if not eq:
        raise HTTPException(status_code=404, detail="Equipment line not found")
    db.delete(eq)
    db.commit()
    return {"ok": True}


@router.get("/{bid_id}/details")
def get_bid_details(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    return {
        "bid_id": bid.id,
        "items": [
            {
                "id": it.id,
                "name": it.name,
                "description": it.description,
                "quantity": it.quantity,
                "unit_cost": it.unit_cost,
                "line_total": round((it.quantity or 0) * (it.unit_cost or 0), 2),
                "supplier_name": it.supplier_name,
                "supplier_lead_time_days": it.supplier_lead_time_days,
                "risk_flag": it.risk_flag,
            }
            for it in bid.items
        ],
        "labor_lines": [
            {
                "id": lb.id,
                "labor_type": lb.labor_type,
                "hourly_rate": lb.hourly_rate,
                "hours": lb.hours,
                "workers": lb.workers,
                "line_total": round((lb.hourly_rate or 0) * (lb.hours or 0) * (lb.workers or 1), 2),
            }
            for lb in bid.labor_lines
        ],
        "transport": (
            {
                "transport_method": bid.transport.transport_method,
                "truck_rental_cost": bid.transport.truck_rental_cost,
                "fuel_cost": bid.transport.fuel_cost,
                "mileage_cost": bid.transport.mileage_cost,
                "toll_fees": bid.transport.toll_fees,
                "driver_cost": bid.transport.driver_cost,
                "trips": bid.transport.trips,
                "delivery_complexity": bid.transport.delivery_complexity,
            }
            if bid.transport
            else None
        ),
        "overhead": (
            {
                "insurance_allocation": bid.overhead.insurance_allocation,
                "storage_cost": bid.overhead.storage_cost,
                "admin_time_cost": bid.overhead.admin_time_cost,
                "bonding_compliance_cost": bid.overhead.bonding_compliance_cost,
                "misc_overhead": bid.overhead.misc_overhead,
            }
            if bid.overhead
            else None
        ),
        "equipment_lines": [
            {
                "id": eq.id,
                "equipment_name": eq.equipment_name,
                "rental_cost": eq.rental_cost,
                "rental_days": eq.rental_days,
                "operator_required": eq.operator_required,
                "operator_cost": eq.operator_cost,
            }
            for eq in bid.equipment_lines
        ],
        "versions": [
            {
                "id": v.id,
                "version_no": v.version_no,
                "selected_mode": v.selected_mode,
                "created_at": v.created_at.isoformat() if v.created_at else None,
                "created_by": v.created_by,
            }
            for v in bid.versions
        ],
    }


# ─── Items CRUD ────────────────────────────────────────────

@router.post("/{bid_id}/items")
def add_item(
    bid_id: int,
    payload: schemas.ItemCreate,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    item = models.BidItem(
        bid_id=bid_id,
        name=payload.name,
        description=payload.description,
        quantity=payload.quantity,
        unit_cost=payload.unit_cost,
        supplier_name=payload.supplier_name,
        supplier_lead_time_days=payload.supplier_lead_time_days,
        risk_flag=payload.risk_flag,
        catalog_item_id=payload.catalog_item_id,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"ok": True, "id": item.id}


@router.delete("/{bid_id}/items/{item_id}")
def delete_item(
    bid_id: int,
    item_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    item = db.query(models.BidItem).filter(
        models.BidItem.id == item_id, models.BidItem.bid_id == bid_id
    ).first()
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")
    db.delete(item)
    db.commit()
    return {"ok": True}


# ─── Labor CRUD ────────────────────────────────────────────

@router.post("/{bid_id}/labor")
def add_labor(
    bid_id: int,
    payload: schemas.LaborCreate,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    labor = models.BidLaborLine(
        bid_id=bid_id,
        labor_type=payload.labor_type,
        hourly_rate=payload.hourly_rate,
        hours=payload.hours,
        workers=payload.workers,
    )
    db.add(labor)
    db.commit()
    db.refresh(labor)
    return {"ok": True, "id": labor.id}


# ─── Transport ─────────────────────────────────────────────

@router.put("/{bid_id}/transport")
def upsert_transport(
    bid_id: int,
    payload: schemas.TransportUpsert,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    if bid.transport:
        for k, v in payload.model_dump().items():
            setattr(bid.transport, k, v)
    else:
        t = models.BidTransport(bid_id=bid_id, **payload.model_dump())
        db.add(t)

    db.commit()
    return {"ok": True}


# ─── Overhead ──────────────────────────────────────────────

@router.put("/{bid_id}/overhead")
def upsert_overhead(
    bid_id: int,
    payload: schemas.OverheadUpsert,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    if bid.overhead:
        for k, v in payload.model_dump().items():
            setattr(bid.overhead, k, v)
    else:
        o = models.BidOverhead(bid_id=bid_id, **payload.model_dump())
        db.add(o)

    db.commit()
    return {"ok": True}


# ─── Equipment ─────────────────────────────────────────────

@router.post("/{bid_id}/equipment")
def add_equipment(
    bid_id: int,
    payload: schemas.EquipmentCreate,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    eq = models.BidEquipmentLine(
        bid_id=bid_id,
        equipment_name=payload.equipment_name,
        rental_cost=payload.rental_cost,
        rental_days=payload.rental_days,
        operator_required=payload.operator_required,
        operator_cost=payload.operator_cost,
    )
    db.add(eq)
    db.commit()
    db.refresh(eq)
    return {"ok": True, "id": eq.id}


# ─── Compute ──────────────────────────────────────────────

@router.post("/{bid_id}/compute")
def compute_bid(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    totals = _compute_totals(bid)
    base_margin = _base_margin(bid)
    recs = _recommendations(totals, base_margin, bid)

    # Check for price drift warnings
    drift_warnings = []
    for it in bid.items:
        if it.catalog_item_id and it.catalog_item:
            cat = it.catalog_item
            if cat.unit_price and it.unit_cost:
                drift = abs(it.unit_cost - cat.unit_price) / max(cat.unit_price, 0.01) * 100
                if drift > 10:
                    drift_warnings.append({
                        "item_id": it.id,
                        "item_name": it.name,
                        "bid_price": it.unit_cost,
                        "catalog_price": cat.unit_price,
                        "drift_pct": round(drift, 1),
                    })

    return {
        "totals": totals,
        "base_margin_pct": base_margin,
        "recommendations": recs,
        "drift_warnings": drift_warnings,
    }


# ─── Approve ──────────────────────────────────────────────

@router.post("/{bid_id}/approve")
def approve_bid(
    bid_id: int,
    payload: schemas.ApproveRequest,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    # Compute fresh totals for the snapshot
    totals = _compute_totals(bid)
    base_margin = _base_margin(bid)
    recs = _recommendations(totals, base_margin, bid)

    selected = next((r for r in recs if r["mode"] == payload.selected_mode), recs[1] if len(recs) > 1 else recs[0])

    # Determine version number
    max_ver = db.query(func.max(models.BidVersion.version_no)).filter(
        models.BidVersion.bid_id == bid_id
    ).scalar() or 0

    snapshot = {
        "totals": totals,
        "base_margin_pct": base_margin,
        "selected": selected,
        "all_recommendations": recs,
        "bid_code": bid.bid_code,
        "contract_title": bid.contract_title,
        "agency_name": bid.agency_name,
        "agency_type": bid.agency_type,
        "contract_type": bid.contract_type,
        "solicitation_number": bid.solicitation_number or "",
        "procurement_method": bid.procurement_method or "",
        "competition_level": bid.competition_level or "",
        "risk_level": bid.risk_level,
        "urgency_level": bid.urgency_level,
        "delivery_distance_miles": bid.delivery_distance_miles or 0,
        "deadline_date": str(bid.deadline_date) if bid.deadline_date else "",
        "desired_profit_mode": bid.desired_profit_mode or "",
        "notes": bid.notes or "",
        "final_bid_price": selected["bid_price"],
        "true_cost": totals["true_cost"],
        "profit_amount": selected["profit_amount"],
        # Line items
        "items": [
            {"name": it.name, "description": it.description or "", "quantity": it.quantity, "unit_cost": float(it.unit_cost or 0), "supplier_name": it.supplier_name or "", "line_total": float((it.quantity or 0) * (it.unit_cost or 0))}
            for it in bid.items
        ],
        # Labor
        "labor_lines": [
            {"labor_type": lb.labor_type, "hourly_rate": float(lb.hourly_rate or 0), "hours": float(lb.hours or 0), "workers": int(lb.workers or 1), "line_total": float((lb.hourly_rate or 0) * (lb.hours or 0) * (lb.workers or 1))}
            for lb in bid.labor_lines
        ],
        # Transport
        "transport": {
            "transport_method": bid.transport.transport_method if bid.transport else "",
            "truck_rental_cost": float(bid.transport.truck_rental_cost or 0) if bid.transport else 0,
            "fuel_cost": float(bid.transport.fuel_cost or 0) if bid.transport else 0,
            "mileage_cost": float(bid.transport.mileage_cost or 0) if bid.transport else 0,
            "toll_fees": float(bid.transport.toll_fees or 0) if bid.transport else 0,
            "driver_cost": float(bid.transport.driver_cost or 0) if bid.transport else 0,
            "trips": int(bid.transport.trips or 1) if bid.transport else 1,
        } if bid.transport else None,
        # Overhead
        "overhead": {
            "insurance_allocation": float(bid.overhead.insurance_allocation or 0) if bid.overhead else 0,
            "storage_cost": float(bid.overhead.storage_cost or 0) if bid.overhead else 0,
            "admin_time_cost": float(bid.overhead.admin_time_cost or 0) if bid.overhead else 0,
            "bonding_compliance_cost": float(bid.overhead.bonding_compliance_cost or 0) if bid.overhead else 0,
            "misc_overhead": float(bid.overhead.misc_overhead or 0) if bid.overhead else 0,
        } if bid.overhead else None,
        # Equipment
        "equipment_lines": [
            {"equipment_name": eq.equipment_name, "rental_cost": float(eq.rental_cost or 0), "rental_days": int(eq.rental_days or 1), "operator_required": eq.operator_required, "operator_cost": float(eq.operator_cost or 0)}
            for eq in bid.equipment_lines
        ],
        "approved_by": payload.approved_by or user,
        "approved_at": datetime.now(timezone.utc).isoformat(),
    }

    justification = payload.assumptions_notes or f"Approved with {payload.selected_mode} mode at {selected['margin_pct']}% margin."

    version = models.BidVersion(
        bid_id=bid_id,
        version_no=max_ver + 1,
        selected_mode=payload.selected_mode,
        totals_json=json.dumps(snapshot),
        justification_text=justification,
        created_by=payload.approved_by or user,
    )
    db.add(version)

    # Update bid status
    bid.status = "approved"
    bid.approved_at = datetime.now(timezone.utc)
    bid.approved_by = payload.approved_by or user

    db.commit()
    db.refresh(version)

    return {
        "ok": True,
        "version_id": version.id,
        "version_no": version.version_no,
        "selected_mode": payload.selected_mode,
        "bid_price": selected["bid_price"],
        "profit": selected["profit_amount"],
    }


# ─── Version / Export ─────────────────────────────────────

@router.get("/versions/{version_id}")
def get_version(
    version_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    ver = db.query(models.BidVersion).filter(models.BidVersion.id == version_id).first()
    if not ver:
        raise HTTPException(status_code=404, detail="Version not found")

    data = json.loads(ver.totals_json)
    data["version_id"] = ver.id
    data["version_no"] = ver.version_no
    data["selected_mode"] = ver.selected_mode
    data["justification_text"] = ver.justification_text
    data["created_at"] = ver.created_at.isoformat() if ver.created_at else None
    data["created_by"] = ver.created_by
    return data


@router.get("/versions/{version_id}/export/csv")
def export_csv(
    version_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    ver = db.query(models.BidVersion).filter(models.BidVersion.id == version_id).first()
    if not ver:
        raise HTTPException(status_code=404, detail="Version not found")

    data = json.loads(ver.totals_json)
    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(["Field", "Value"])
    writer.writerow(["Bid Code", data.get("bid_code", "")])
    writer.writerow(["Contract", data.get("contract_title", "")])
    writer.writerow(["Agency", data.get("agency_name", "")])
    writer.writerow(["Version", ver.version_no])
    writer.writerow(["Mode", ver.selected_mode])
    writer.writerow([])

    totals = data.get("totals", {})
    writer.writerow(["Cost Category", "Amount"])
    for k, v in totals.items():
        writer.writerow([k.replace("_", " ").title(), f"${v:,.2f}" if isinstance(v, (int, float)) else str(v)])
    writer.writerow([])

    selected = data.get("selected", {})
    writer.writerow(["Final Bid Price", f"${selected.get('bid_price', 0):,.2f}"])
    writer.writerow(["Profit", f"${selected.get('profit_amount', 0):,.2f}"])
    writer.writerow(["Margin", f"{selected.get('margin_pct', 0):.2f}%"])

    output.seek(0)
    code = data.get("bid_code", f"version-{version_id}")
    return StreamingResponse(
        output,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={code}-proposal.csv"},
    )


@router.get("/versions/{version_id}/export/pdf")
def export_pdf(
    version_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    ver = db.query(models.BidVersion).filter(models.BidVersion.id == version_id).first()
    if not ver:
        raise HTTPException(status_code=404, detail="Version not found")

    data = json.loads(ver.totals_json)

    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak, KeepTogether,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )

    # ─── Colors ───────────────────────────────────────────
    DARK = colors.HexColor("#0B1020")
    PURPLE = colors.HexColor("#7A3FFF")
    GOLD = colors.HexColor("#D7B66D")
    LIGHT_BG = colors.HexColor("#F4F1EB")
    WHITE = colors.white
    GREY = colors.HexColor("#555555")
    LIGHT_GREY = colors.HexColor("#E8E8E8")

    # ─── Styles ───────────────────────────────────────────
    styles = getSampleStyleSheet()
    s_title = ParagraphStyle("s_title", parent=styles["Title"], fontSize=24, leading=28, textColor=DARK, fontName="Helvetica-Bold", spaceAfter=4)
    s_subtitle = ParagraphStyle("s_subtitle", parent=styles["Normal"], fontSize=11, textColor=GREY, spaceAfter=12)
    s_h2 = ParagraphStyle("s_h2", parent=styles["Heading2"], fontSize=14, leading=18, textColor=PURPLE, fontName="Helvetica-Bold", spaceBefore=16, spaceAfter=8)
    s_h3 = ParagraphStyle("s_h3", parent=styles["Heading3"], fontSize=11, leading=14, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    s_body = ParagraphStyle("s_body", parent=styles["Normal"], fontSize=10, leading=14, textColor=DARK)
    s_small = ParagraphStyle("s_small", parent=styles["Normal"], fontSize=8.5, leading=11, textColor=GREY)
    s_right = ParagraphStyle("s_right", parent=s_body, alignment=TA_RIGHT)
    s_center = ParagraphStyle("s_center", parent=s_body, alignment=TA_CENTER)
    s_gold = ParagraphStyle("s_gold", parent=s_body, fontSize=18, fontName="Helvetica-Bold", textColor=GOLD, alignment=TA_CENTER)
    s_big_num = ParagraphStyle("s_big_num", parent=s_body, fontSize=22, fontName="Helvetica-Bold", textColor=DARK, alignment=TA_CENTER)

    def money(x):
        try:
            return f"${float(x):,.2f}"
        except (TypeError, ValueError):
            return "$0.00"

    def pct(x):
        try:
            return f"{float(x):.2f}%"
        except (TypeError, ValueError):
            return "0.00%"

    story = []
    totals = data.get("totals", {})
    selected = data.get("selected", {})
    bid_code = data.get("bid_code", "")
    items = data.get("items", [])
    labor_lines = data.get("labor_lines", [])
    transport = data.get("transport")
    overhead_data = data.get("overhead")
    equipment_lines = data.get("equipment_lines", [])
    all_recs = data.get("all_recommendations", [])

    # ═══════════════════════════════════════════════════════
    # PAGE 1: COVER / EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════

    # Header band
    header_data = [[
        Paragraph("SENTRIBID", ParagraphStyle("hdr", parent=s_body, fontSize=10, fontName="Helvetica-Bold", textColor=WHITE)),
        Paragraph("BID PROPOSAL", ParagraphStyle("hdr2", parent=s_body, fontSize=10, fontName="Helvetica-Bold", textColor=WHITE, alignment=TA_RIGHT)),
    ]]
    header_table = Table(header_data, colWidths=[3.5 * inch, 3.5 * inch])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), DARK),
        ("TEXTCOLOR", (0, 0), (-1, -1), WHITE),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("ROUNDEDCORNERS", [6, 6, 6, 6]),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 20))

    # Title
    story.append(Paragraph(data.get("contract_title", "Bid Proposal"), s_title))
    story.append(Paragraph(f"Prepared for <b>{data.get('agency_name', '')}</b> &nbsp;|&nbsp; {data.get('agency_type', '').title()} Agency &nbsp;|&nbsp; {bid_code}", s_subtitle))
    story.append(HRFlowable(width="100%", thickness=1.5, color=GOLD, spaceAfter=16))

    # Key Metrics Cards
    bid_price = selected.get("bid_price", 0)
    profit_amt = selected.get("profit_amount", 0)
    margin_val = selected.get("margin_pct", 0)

    metrics = [[
        Paragraph("PROPOSED BID PRICE", s_center),
        Paragraph("ESTIMATED PROFIT", s_center),
        Paragraph("PROFIT MARGIN", s_center),
        Paragraph("WIN SCORE", s_center),
    ], [
        Paragraph(money(bid_price), s_big_num),
        Paragraph(money(profit_amt), s_big_num),
        Paragraph(pct(margin_val), s_big_num),
        Paragraph(str(selected.get("win_score", "—")), s_big_num),
    ]]
    metrics_table = Table(metrics, colWidths=[1.75 * inch] * 4)
    metrics_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BG),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, 0), 10),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 12),
        ("BOX", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
        ("LINEBELOW", (0, 0), (-1, 0), 0.5, LIGHT_GREY),
        ("LINEBEFORE", (1, 0), (1, -1), 0.5, LIGHT_GREY),
        ("LINEBEFORE", (2, 0), (2, -1), 0.5, LIGHT_GREY),
        ("LINEBEFORE", (3, 0), (3, -1), 0.5, LIGHT_GREY),
        ("ROUNDEDCORNERS", [6, 6, 6, 6]),
    ]))
    story.append(metrics_table)
    story.append(Spacer(1, 16))

    # Bid Info Table
    story.append(Paragraph("Bid Information", s_h2))
    info_rows = [
        ["Bid Code", bid_code, "Pricing Mode", (ver.selected_mode or "").title()],
        ["Contract Type", data.get("contract_type", "").title(), "Competition", data.get("competition_level", "").title()],
        ["Solicitation #", data.get("solicitation_number", "") or "N/A", "Procurement", (data.get("procurement_method", "") or "").upper()],
        ["Deadline", data.get("deadline_date", ""), "Risk Level", f"{data.get('risk_level', '')} / 5"],
        ["Delivery Distance", f"{data.get('delivery_distance_miles', 0)} miles", "Urgency", f"{data.get('urgency_level', '')} / 5"],
        ["Version", str(ver.version_no), "Approved By", data.get("approved_by", ver.created_by or "")],
    ]
    info_data = []
    for row in info_rows:
        info_data.append([
            Paragraph(f"<b>{row[0]}</b>", s_small), Paragraph(str(row[1]), s_body),
            Paragraph(f"<b>{row[2]}</b>", s_small), Paragraph(str(row[3]), s_body),
        ])
    info_table = Table(info_data, colWidths=[1.2 * inch, 2.3 * inch, 1.2 * inch, 2.3 * inch])
    info_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F8F6F2")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F8F6F2")),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(info_table)

    # Justification / Notes
    if ver.justification_text:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Justification / Notes", s_h3))
        story.append(Paragraph(ver.justification_text, s_body))

    if data.get("notes"):
        story.append(Spacer(1, 6))
        story.append(Paragraph(f"<i>{data['notes']}</i>", s_small))

    # ═══════════════════════════════════════════════════════
    # PAGE 2: COST BREAKDOWN
    # ═══════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("Cost Breakdown", s_h2))
    story.append(Paragraph("Detailed breakdown of all cost components that make up the true cost of this bid.", s_small))
    story.append(Spacer(1, 8))

    # Cost Summary Table
    cost_rows = [
        [Paragraph("<b>Category</b>", s_small), Paragraph("<b>Amount</b>", ParagraphStyle("rh", parent=s_small, alignment=TA_RIGHT)), Paragraph("<b>% of Total</b>", ParagraphStyle("rh2", parent=s_small, alignment=TA_RIGHT))],
    ]
    true_cost = totals.get("true_cost", 0) or 1
    cost_items = [
        ("Materials / Supplies", totals.get("item_subtotal", 0)),
        ("Labor", totals.get("labor_total", 0)),
        ("Transport / Delivery", totals.get("transport_total", 0)),
        ("Equipment Rentals", totals.get("equipment_total", 0)),
        ("Overhead / Indirect", totals.get("overhead_total", 0)),
    ]
    for label, amt in cost_items:
        pct_val = (amt / true_cost * 100) if true_cost else 0
        cost_rows.append([
            Paragraph(label, s_body),
            Paragraph(money(amt), s_right),
            Paragraph(f"{pct_val:.1f}%", s_right),
        ])

    # Totals
    cost_rows.append([Paragraph("<b>True Cost</b>", ParagraphStyle("tc", parent=s_body, fontName="Helvetica-Bold")), Paragraph(f"<b>{money(true_cost)}</b>", ParagraphStyle("tcr", parent=s_right, fontName="Helvetica-Bold")), Paragraph("100.0%", s_right)])
    cost_rows.append([Paragraph("Risk Buffer", s_body), Paragraph(money(totals.get("risk_buffer", 0)), s_right), Paragraph("", s_right)])
    cost_rows.append([Paragraph("<b>Adjusted Cost</b>", ParagraphStyle("ac", parent=s_body, fontName="Helvetica-Bold")), Paragraph(f"<b>{money(totals.get('adjusted_cost', 0))}</b>", ParagraphStyle("acr", parent=s_right, fontName="Helvetica-Bold")), Paragraph("", s_right)])

    cost_table = Table(cost_rows, colWidths=[3.2 * inch, 2.0 * inch, 1.8 * inch])
    cost_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
        ("BACKGROUND", (0, 0), (-1, 0), DARK),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("BACKGROUND", (0, 6), (-1, 6), colors.HexColor("#F0EDE6")),
        ("BACKGROUND", (0, 8), (-1, 8), colors.HexColor("#EDE8DF")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("LINEABOVE", (0, 6), (-1, 6), 1, DARK),
    ]))
    story.append(cost_table)

    # ─── Line Items Detail ────────────────────────────────
    if items:
        story.append(Spacer(1, 16))
        story.append(Paragraph("Line Items Detail", s_h2))
        item_rows = [
            [Paragraph("<b>#</b>", s_small), Paragraph("<b>Item</b>", s_small), Paragraph("<b>Qty</b>", ParagraphStyle("qh", parent=s_small, alignment=TA_CENTER)), Paragraph("<b>Unit Cost</b>", ParagraphStyle("uh", parent=s_small, alignment=TA_RIGHT)), Paragraph("<b>Total</b>", ParagraphStyle("th", parent=s_small, alignment=TA_RIGHT)), Paragraph("<b>Supplier</b>", s_small)],
        ]
        for idx, it in enumerate(items, 1):
            desc_text = it.get("name", "")
            if it.get("description"):
                desc_text += f"<br/><font size=8 color='#777'>{it['description']}</font>"
            item_rows.append([
                Paragraph(str(idx), s_center),
                Paragraph(desc_text, s_body),
                Paragraph(str(it.get("quantity", 0)), s_center),
                Paragraph(money(it.get("unit_cost", 0)), s_right),
                Paragraph(money(it.get("line_total", 0)), s_right),
                Paragraph(it.get("supplier_name", "—"), s_small),
            ])
        # Subtotal row
        item_rows.append([
            Paragraph("", s_body), Paragraph("<b>Subtotal</b>", ParagraphStyle("st", parent=s_body, fontName="Helvetica-Bold")),
            Paragraph("", s_body), Paragraph("", s_body),
            Paragraph(f"<b>{money(totals.get('item_subtotal', 0))}</b>", ParagraphStyle("str", parent=s_right, fontName="Helvetica-Bold")),
            Paragraph("", s_body),
        ])

        item_table = Table(item_rows, colWidths=[0.4 * inch, 2.3 * inch, 0.6 * inch, 1.1 * inch, 1.1 * inch, 1.5 * inch])
        item_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
            ("BACKGROUND", (0, 0), (-1, 0), DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F0EDE6")),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(item_table)

    # ─── Labor Detail ─────────────────────────────────────
    if labor_lines:
        story.append(Spacer(1, 16))
        story.append(Paragraph("Labor Breakdown", s_h2))
        labor_rows = [
            [Paragraph("<b>Role</b>", s_small), Paragraph("<b>Rate</b>", ParagraphStyle("lr", parent=s_small, alignment=TA_RIGHT)), Paragraph("<b>Hours</b>", ParagraphStyle("lh", parent=s_small, alignment=TA_CENTER)), Paragraph("<b>Workers</b>", ParagraphStyle("lw", parent=s_small, alignment=TA_CENTER)), Paragraph("<b>Total</b>", ParagraphStyle("lt", parent=s_small, alignment=TA_RIGHT))],
        ]
        for lb in labor_lines:
            labor_rows.append([
                Paragraph(lb.get("labor_type", ""), s_body),
                Paragraph(money(lb.get("hourly_rate", 0)) + "/hr", s_right),
                Paragraph(str(lb.get("hours", 0)), s_center),
                Paragraph(str(lb.get("workers", 1)), s_center),
                Paragraph(money(lb.get("line_total", 0)), s_right),
            ])
        labor_rows.append([
            Paragraph("<b>Total</b>", ParagraphStyle("ltt", parent=s_body, fontName="Helvetica-Bold")),
            Paragraph("", s_body), Paragraph("", s_body), Paragraph("", s_body),
            Paragraph(f"<b>{money(totals.get('labor_total', 0))}</b>", ParagraphStyle("lttr", parent=s_right, fontName="Helvetica-Bold")),
        ])
        labor_table = Table(labor_rows, colWidths=[2.2 * inch, 1.3 * inch, 1.0 * inch, 1.0 * inch, 1.5 * inch])
        labor_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
            ("BACKGROUND", (0, 0), (-1, 0), DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F0EDE6")),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(labor_table)

    # ─── Transport & Overhead (side by side if short) ─────
    if transport and any(v for k, v in transport.items() if k != "transport_method" and v):
        story.append(Spacer(1, 16))
        story.append(Paragraph("Transport / Delivery", s_h3))
        t = transport
        per_trip = t.get("truck_rental_cost", 0) + t.get("fuel_cost", 0) + t.get("mileage_cost", 0) + t.get("toll_fees", 0) + t.get("driver_cost", 0)
        trans_rows = [
            ["Method", (t.get("transport_method", "") or "").title()],
            ["Vehicle Rental", money(t.get("truck_rental_cost", 0))],
            ["Fuel", money(t.get("fuel_cost", 0))],
            ["Mileage", money(t.get("mileage_cost", 0))],
            ["Tolls", money(t.get("toll_fees", 0))],
            ["Driver", money(t.get("driver_cost", 0))],
            ["Per-Trip Cost", money(per_trip)],
            ["Trips", str(t.get("trips", 1))],
            ["Total", money(per_trip * t.get("trips", 1))],
        ]
        trans_t = Table(trans_rows, colWidths=[2.0 * inch, 2.0 * inch])
        trans_t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F0EDE6")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(trans_t)

    if overhead_data and any(v for v in overhead_data.values() if v):
        story.append(Spacer(1, 12))
        story.append(Paragraph("Overhead / Indirect Costs", s_h3))
        ov = overhead_data
        ov_rows = [
            ["Insurance", money(ov.get("insurance_allocation", 0))],
            ["Storage", money(ov.get("storage_cost", 0))],
            ["Admin Time", money(ov.get("admin_time_cost", 0))],
            ["Bonding / Compliance", money(ov.get("bonding_compliance_cost", 0))],
            ["Misc", money(ov.get("misc_overhead", 0))],
            ["Total", money(totals.get("overhead_total", 0))],
        ]
        ov_t = Table(ov_rows, colWidths=[2.5 * inch, 2.0 * inch])
        ov_t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F0EDE6")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(ov_t)

    # ═══════════════════════════════════════════════════════
    # PAGE 3: PRICING ANALYSIS
    # ═══════════════════════════════════════════════════════
    if all_recs:
        story.append(PageBreak())
        story.append(Paragraph("Pricing Analysis", s_h2))
        story.append(Paragraph("Three pricing strategies were evaluated. The selected mode is highlighted.", s_small))
        story.append(Spacer(1, 8))

        rec_header = [
            Paragraph("<b>Mode</b>", ParagraphStyle("mh", parent=s_small, textColor=WHITE)),
            Paragraph("<b>Bid Price</b>", ParagraphStyle("mh2", parent=s_small, textColor=WHITE, alignment=TA_RIGHT)),
            Paragraph("<b>Profit</b>", ParagraphStyle("mh3", parent=s_small, textColor=WHITE, alignment=TA_RIGHT)),
            Paragraph("<b>Margin</b>", ParagraphStyle("mh4", parent=s_small, textColor=WHITE, alignment=TA_RIGHT)),
            Paragraph("<b>Win Score</b>", ParagraphStyle("mh5", parent=s_small, textColor=WHITE, alignment=TA_CENTER)),
        ]
        rec_rows = [rec_header]

        for r in all_recs:
            is_selected = r.get("mode", "") == ver.selected_mode
            mode_label = r.get("mode", "").upper()
            if is_selected:
                mode_label += " (SELECTED)"
            rec_rows.append([
                Paragraph(f"<b>{mode_label}</b>" if is_selected else mode_label, s_body),
                Paragraph(money(r.get("bid_price", 0)), s_right),
                Paragraph(money(r.get("profit_amount", 0)), s_right),
                Paragraph(pct(r.get("margin_pct", 0)), s_right),
                Paragraph(str(r.get("win_score", "—")), s_center),
            ])

        rec_table = Table(rec_rows, colWidths=[1.8 * inch, 1.6 * inch, 1.6 * inch, 1.0 * inch, 1.0 * inch])
        highlight_rows = []
        for idx, r in enumerate(all_recs):
            if r.get("mode", "") == ver.selected_mode:
                highlight_rows.append(idx + 1)

        rec_style = [
            ("GRID", (0, 0), (-1, -1), 0.5, LIGHT_GREY),
            ("BACKGROUND", (0, 0), (-1, 0), DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]
        for hr in highlight_rows:
            rec_style.append(("BACKGROUND", (0, hr), (-1, hr), colors.HexColor("#F5F0E4")))
            rec_style.append(("BOX", (0, hr), (-1, hr), 1.5, GOLD))

        rec_table.setStyle(TableStyle(rec_style))
        story.append(rec_table)

        # Warnings
        for r in all_recs:
            warnings = r.get("warnings", [])
            if warnings and r.get("mode", "") == ver.selected_mode:
                story.append(Spacer(1, 8))
                for w in warnings:
                    story.append(Paragraph(f"<font color='#CC6600'>&#9888; {w}</font>", s_small))

    # ─── Footer: Confidentiality ──────────────────────────
    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT_GREY))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        f"<font size=8 color='#999'>This document was generated by SentriBiD on {datetime.now(timezone.utc).strftime('%B %d, %Y at %H:%M UTC')}. "
        f"Bid {bid_code}, Version {ver.version_no}. Confidential — for internal use and authorized recipients only.</font>",
        s_small,
    ))

    # Build
    doc.build(story)
    buf.seek(0)

    code = data.get("bid_code", f"version-{version_id}")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={code}-proposal.pdf"},
    )


@router.get("/versions/{version_id}/export/docx")
def export_docx(
    version_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    ver = db.query(models.BidVersion).filter(models.BidVersion.id == version_id).first()
    if not ver:
        raise HTTPException(status_code=404, detail="Version not found")

    data = json.loads(ver.totals_json)

    from docx import Document

    doc = Document()
    doc.add_heading("SentriBiD — Bid Proposal", 0)

    doc.add_paragraph(f"Bid Code: {data.get('bid_code', '')}")
    doc.add_paragraph(f"Contract: {data.get('contract_title', '')}")
    doc.add_paragraph(f"Agency: {data.get('agency_name', '')}")
    doc.add_paragraph(f"Version: {ver.version_no}  |  Mode: {ver.selected_mode}")

    doc.add_heading("Cost Breakdown", level=1)
    totals = data.get("totals", {})
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Amount"
    for k, v in totals.items():
        row = table.add_row().cells
        row[0].text = k.replace("_", " ").title()
        row[1].text = f"${v:,.2f}" if isinstance(v, (int, float)) else str(v)

    doc.add_heading("Final Pricing", level=1)
    selected = data.get("selected", {})
    doc.add_paragraph(f"Final Bid Price: ${selected.get('bid_price', 0):,.2f}")
    doc.add_paragraph(f"Profit: ${selected.get('profit_amount', 0):,.2f}")
    doc.add_paragraph(f"Margin: {selected.get('margin_pct', 0):.2f}%")

    doc.add_heading("Justification", level=1)
    doc.add_paragraph(ver.justification_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    code = data.get("bid_code", f"version-{version_id}")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={code}-proposal.docx"},
    )


# ─── AI Copilot Endpoints ─────────────────────────────────

@router.post("/{bid_id}/copilot/analyze")
def copilot_full_analysis(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Run full AI copilot analysis: risk + profit + compliance."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    totals = _compute_totals(bid)
    base_margin = _base_margin(bid)
    recs = _recommendations(totals, base_margin, bid)

    context = {
        "bid": {
            "id": bid.id,
            "bid_code": bid.bid_code,
            "contract_title": bid.contract_title,
            "agency_name": bid.agency_name,
            "agency_type": bid.agency_type,
            "contract_type": bid.contract_type,
            "competition_level": bid.competition_level,
            "risk_level": bid.risk_level,
            "urgency_level": bid.urgency_level,
            "delivery_distance_miles": bid.delivery_distance_miles,
            "deadline_date": str(bid.deadline_date) if bid.deadline_date else None,
            "desired_profit_mode": bid.desired_profit_mode,
            "min_acceptable_profit": bid.min_acceptable_profit,
            "margin_override_pct": bid.margin_override_pct,
            "status": bid.status,
            "notes": bid.notes,
            "item_count": len(bid.items),
            "labor_count": len(bid.labor_lines),
        },
        "totals": totals,
        "base_margin_pct": base_margin,
        "recommendations": recs,
        "attachments_text": _get_attachments_text(bid),
    }

    risk = analyze_bid_risk(context)
    profit = analyze_bid_profit(context)
    compliance = analyze_bid_compliance(context)

    now = datetime.now(timezone.utc)

    # Store on bid
    bid.ai_risk_analysis = json.dumps(risk)
    bid.ai_profit_suggestions = json.dumps(profit)
    bid.ai_compliance_flags = json.dumps(compliance)
    bid.ai_analyzed_at = now
    db.commit()

    executive_summary = (
        f"Bid {bid.bid_code} for {bid.agency_name}: "
        f"Risk grade {risk['risk_grade']} (score {risk['overall_risk_score']}/100). "
        f"{compliance['overall_status'].replace('_', ' ').title()} on compliance. "
        f"Recommended mode: {profit['optimal_mode']}. "
        f"{len(profit['suggestions'])} optimization opportunities identified."
    )

    return {
        "bid_id": bid.id,
        "bid_code": bid.bid_code,
        "risk": risk,
        "profit": profit,
        "compliance": compliance,
        "executive_summary": executive_summary,
        "analyzed_at": now.isoformat(),
    }


@router.post("/{bid_id}/copilot/chat")
def copilot_chat(
    bid_id: int,
    payload: schemas.CopilotChatRequest,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Chat with AI copilot about a specific bid."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    totals = _compute_totals(bid)

    context = {
        "bid_code": bid.bid_code,
        "contract_title": bid.contract_title,
        "agency_name": bid.agency_name,
        "agency_type": bid.agency_type,
        "contract_type": bid.contract_type,
        "competition_level": bid.competition_level,
        "risk_level": bid.risk_level,
        "status": bid.status,
        "totals": totals,
        "notes": bid.notes,
        "attachments_text": _get_attachments_text(bid),
    }

    result = chat_with_copilot(payload.message, context)
    return result


# ═══════════════════════════════════════════════════════════
# ATTACHMENTS
# ═══════════════════════════════════════════════════════════

def _get_attachments_text(bid) -> str:
    """Gather extracted text from all bid attachments for AI context."""
    texts = []
    for att in bid.attachments:
        if att.extracted_text:
            texts.append(f"[{att.category.upper()}: {att.filename}]\n{att.extracted_text[:1500]}")
    return "\n\n".join(texts)[:4000] if texts else ""


def _extract_pdf_text(filepath: str) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:20]:  # Limit to 20 pages
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)[:5000]
    except Exception as e:
        return f"[PDF text extraction failed: {e}]"


def _extract_docx_text(filepath: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:5000]
    except Exception as e:
        return f"[DOCX text extraction failed: {e}]"


def _extract_text(filepath: str, file_type: str) -> str:
    """Route text extraction by file type."""
    if file_type == "pdf":
        return _extract_pdf_text(filepath)
    elif file_type in ("docx", "doc"):
        return _extract_docx_text(filepath)
    elif file_type in ("txt", "csv", "rtf"):
        try:
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                return f.read()[:5000]
        except Exception:
            return ""
    return ""


@router.post("/{bid_id}/attachments")
async def upload_attachment(
    bid_id: int,
    file: UploadFile = File(...),
    category: str = Form("general"),
    description: str = Form(""),
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Upload a file attachment to a bid."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    # Validate extension
    ext = (file.filename or "").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type '.{ext}' not allowed. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    # Read and check size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB")

    # Save to disk
    bid_dir = UPLOAD_DIR / str(bid_id)
    bid_dir.mkdir(parents=True, exist_ok=True)

    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    stored_path = bid_dir / safe_name
    with open(stored_path, "wb") as f:
        f.write(content)

    # Extract text for AI context
    extracted_text = _extract_text(str(stored_path), ext)

    # Save to DB
    attachment = models.BidAttachment(
        bid_id=bid_id,
        filename=file.filename or "unknown",
        stored_path=str(stored_path),
        file_type=ext,
        file_size=len(content),
        category=category,
        description=description,
        extracted_text=extracted_text if extracted_text else None,
        uploaded_by=user,
    )
    db.add(attachment)
    db.commit()
    db.refresh(attachment)

    return {
        "id": attachment.id,
        "filename": attachment.filename,
        "file_type": ext,
        "file_size": len(content),
        "category": category,
        "description": description,
        "has_extracted_text": bool(extracted_text),
        "uploaded_at": attachment.uploaded_at.isoformat() if attachment.uploaded_at else None,
    }


@router.get("/{bid_id}/attachments")
def list_attachments(
    bid_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """List all attachments for a bid."""
    bid = db.query(models.Bid).filter(models.Bid.id == bid_id).first()
    if not bid:
        raise HTTPException(status_code=404, detail="Bid not found")

    return [
        {
            "id": a.id,
            "filename": a.filename,
            "file_type": a.file_type,
            "file_size": a.file_size,
            "category": a.category,
            "description": a.description,
            "has_extracted_text": bool(a.extracted_text),
            "uploaded_at": a.uploaded_at.isoformat() if a.uploaded_at else None,
            "uploaded_by": a.uploaded_by,
        }
        for a in bid.attachments
    ]


@router.get("/{bid_id}/attachments/{attachment_id}/download")
def download_attachment(
    bid_id: int,
    attachment_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Download an attachment file."""
    att = db.query(models.BidAttachment).filter(
        models.BidAttachment.id == attachment_id,
        models.BidAttachment.bid_id == bid_id,
    ).first()
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")

    if not os.path.exists(att.stored_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    return FileResponse(
        att.stored_path,
        filename=att.filename,
        media_type="application/octet-stream",
    )


@router.delete("/{bid_id}/attachments/{attachment_id}")
def delete_attachment(
    bid_id: int,
    attachment_id: int,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Delete an attachment."""
    att = db.query(models.BidAttachment).filter(
        models.BidAttachment.id == attachment_id,
        models.BidAttachment.bid_id == bid_id,
    ).first()
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")

    # Remove file from disk
    try:
        if os.path.exists(att.stored_path):
            os.remove(att.stored_path)
    except Exception:
        pass

    db.delete(att)
    db.commit()
    return {"ok": True, "deleted_id": attachment_id}


