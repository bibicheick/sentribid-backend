# backend/app/routers/opportunities.py
"""
Opportunity Discovery & Intelligent Contract Explainer
Phase 0: Manual entry + upload + AI analysis + convert to bid + proposal generation
"""

import os, json, uuid, shutil, logging
from pathlib import Path
from datetime import datetime, date, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc

from ..db import get_db
from ..auth import require_auth
from .. import models, schemas

logger = logging.getLogger("sentribid.opportunities")

router = APIRouter(prefix="/opportunities", tags=["opportunities"])

UPLOAD_DIR = os.getenv("SENTRIBID_UPLOAD_DIR", "./uploads/opportunities")
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "xlsx", "xls", "csv", "txt", "rtf"}


# ─── Helpers ──────────────────────────────────────────────

def _gen_opp_code():
    short = uuid.uuid4().hex[:8].upper()
    return f"OP-{datetime.now().year}-{short}"


def _extract_text(filepath: str, file_type: str) -> str:
    text = ""
    try:
        if file_type == "pdf":
            try:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    for i, page in enumerate(pdf.pages[:50]):
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                        if len(text) > 15000:
                            break
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        elif file_type in ("docx", "doc"):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                    if len(text) > 15000:
                        break
                # Also extract tables from DOCX
                for table in doc.tables[:20]:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text += row_text + "\n"
                    if len(text) > 15000:
                        break
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
        elif file_type in ("xlsx", "xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
                for sheet_name in wb.sheetnames[:10]:
                    ws = wb[sheet_name]
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    row_count = 0
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c).strip() if c is not None else "" for c in row]
                        if any(cells):
                            text += " | ".join(cells) + "\n"
                            row_count += 1
                        if row_count > 500 or len(text) > 15000:
                            break
                    if len(text) > 15000:
                        break
                wb.close()
            except Exception as e:
                logger.warning(f"Excel extraction failed: {e}")
                # Fallback: try csv-style read
                try:
                    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                        text = f.read(15000)
                except Exception:
                    pass
        elif file_type == "csv":
            try:
                import csv
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    reader = csv.reader(f)
                    for i, row in enumerate(reader):
                        text += " | ".join(row) + "\n"
                        if i > 500 or len(text) > 15000:
                            break
            except Exception as e:
                logger.warning(f"CSV extraction failed: {e}")
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read(15000)
        elif file_type in ("txt", "rtf"):
            with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                text = f.read(15000)
    except Exception as e:
        logger.warning(f"Text extraction failed: {e}")
    return text[:15000]


def _get_attachments_text(opp) -> str:
    parts = []
    for att in (opp.attachments or []):
        if att.extracted_text:
            parts.append(f"[{att.filename}]\n{att.extracted_text}")
    combined = "\n\n".join(parts)
    return combined[:25000]


def _call_openai(system: str, user: str, json_mode: bool = True):
    try:
        api_key = os.getenv("OPENAI_API_KEY", "")
        if not api_key:
            return None
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        kwargs = {
            "model": model,
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
            "temperature": 0.2,
            "max_tokens": 2000,
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message.content
    except Exception as e:
        logger.error(f"OpenAI call failed: {e}")
        return None


def _build_opp_context(opp, attachments_text: str = "") -> str:
    parts = [
        f"Title: {opp.title}",
        f"Agency: {opp.agency_name} ({opp.agency_type or 'unknown'})",
        f"Solicitation: {opp.solicitation_number or 'N/A'}",
        f"NAICS: {opp.naics_code or 'N/A'}",
        f"Set-Aside: {opp.set_aside or 'N/A'}",
        f"Contract Type: {opp.contract_type or 'N/A'}",
        f"Value Range: ${opp.estimated_value_low or '?'} - ${opp.estimated_value_high or '?'}",
        f"Location: {opp.location_city or ''}, {opp.location_state or ''}",
        f"Posted: {opp.posted_date or 'N/A'}",
        f"Due: {opp.due_date or 'N/A'}",
        f"Source: {opp.source_url or 'N/A'}",
    ]
    if opp.description:
        parts.append(f"\nDescription:\n{opp.description[:3000]}")
    if attachments_text:
        parts.append(f"\n--- Solicitation Documents ---\n{attachments_text}")
    return "\n".join(parts)


def _get_business_profile(db: Session, user_email: str) -> Optional[str]:
    """Get the user's business profile text for AI context."""
    user = db.query(models.User).filter(models.User.email == user_email).first()
    if not user or not user.profile:
        return None
    p = user.profile
    parts = []
    if p.company_name:
        parts.append(f"Company: {p.company_name}")
    if p.company_description:
        parts.append(f"Description: {p.company_description}")
    if p.elevator_pitch:
        parts.append(f"Pitch: {p.elevator_pitch}")
    if p.naics_codes:
        parts.append(f"NAICS: {p.naics_codes}")
    if p.certifications:
        parts.append(f"Certifications: {p.certifications}")
    if p.core_competencies:
        parts.append(f"Core Competencies: {p.core_competencies}")
    if p.differentiators:
        parts.append(f"Differentiators: {p.differentiators}")
    if p.past_performance:
        parts.append(f"Past Performance: {p.past_performance[:1000]}")
    if p.key_personnel:
        parts.append(f"Key Personnel: {p.key_personnel[:500]}")
    if p.capability_statement_text:
        parts.append(f"Capability Statement: {p.capability_statement_text[:1500]}")
    if p.company_size:
        parts.append(f"Size: {p.company_size}")
    if p.annual_revenue:
        parts.append(f"Revenue: {p.annual_revenue}")
    return "\n".join(parts) if parts else None


# ─── AI Analysis Functions ────────────────────────────────

def _ai_executive_summary(context: str) -> Optional[str]:
    system = """You are a government contracting analyst. Summarize this solicitation in plain English.
Return JSON: {"summary": "...", "procurement_type": "...", "key_dates": {"posted": "...", "due": "..."}, "value_range": "...", "critical_insight": "..."}"""
    return _call_openai(system, context)


def _ai_requirements(context: str) -> Optional[str]:
    system = """You are a government contracting analyst. Extract ALL requirements from this solicitation.
Categorize each as: mandatory, desirable, technical, personnel, past_performance, or compliance.
Return JSON: {"mandatory": [{"requirement": "...", "source": "...", "confidence": "high/medium/low"}], "desirable": [...], "technical": [...], "personnel": [...], "past_performance": [...], "compliance": [...]}"""
    return _call_openai(system, context)


def _ai_evaluation_factors(context: str) -> Optional[str]:
    system = """You are a government contracting analyst. Identify evaluation criteria and their weights.
Return JSON: {"factors": [{"name": "...", "weight_pct": 0, "description": "...", "strong_response_guidance": "..."}], "scoring_method": "..."}"""
    return _call_openai(system, context)


def _ai_risk_flags(context: str) -> Optional[str]:
    system = """You are a government contracting risk analyst. Identify risks across: timeline, technical, compliance, competitive, financial.
Return JSON: {"risks": [{"category": "...", "severity": "high/medium/low", "title": "...", "description": "...", "mitigation": "..."}], "overall_risk_level": "high/medium/low"}"""
    return _call_openai(system, context)


def _ai_compliance_checklist(context: str) -> Optional[str]:
    system = """You are a government contracting compliance specialist. Generate a compliance checklist.
Return JSON: {"items": [{"requirement": "...", "category": "...", "urgency": "high/medium/low", "action_needed": "..."}], "total_items": 0}"""
    return _call_openai(system, context)


def _ai_bid_strategy(context: str) -> Optional[str]:
    system = """You are a government bid strategist. Recommend positioning strategy.
Return JSON: {"pricing_approach": "aggressive/balanced/conservative", "pricing_reasoning": "...", "differentiators": ["..."], "win_themes": ["..."], "teaming_recommendation": "...", "competitive_positioning": "..."}"""
    return _call_openai(system, context)


def _ai_bid_recommendation(context: str, prior_analysis: str) -> Optional[str]:
    system = """You are a senior government contracting advisor. Given the opportunity and prior analysis, provide a bid/no-bid recommendation.
Return JSON: {"recommendation": "bid/no_bid/conditional", "confidence": 0.0-1.0, "go_factors": ["..."], "no_go_factors": ["..."], "conditions": ["..."], "executive_reasoning": "...", "suggested_bid_price_low": 0, "suggested_bid_price_high": 0, "estimated_profit_margin_pct": 0}"""
    full_context = f"{context}\n\n--- Prior Analysis ---\n{prior_analysis}"
    return _call_openai(system, full_context)


# ─── CRUD Endpoints ───────────────────────────────────────

@router.get("")
def list_opportunities(
    status: Optional[str] = None,
    search: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    q = db.query(models.Opportunity).order_by(desc(models.Opportunity.id))
    if status:
        q = q.filter(models.Opportunity.status == status)
    if search:
        like = f"%{search}%"
        from sqlalchemy import or_
        q = q.filter(or_(
            models.Opportunity.title.ilike(like),
            models.Opportunity.agency_name.ilike(like),
            models.Opportunity.naics_code.ilike(like),
            models.Opportunity.solicitation_number.ilike(like),
        ))
    return q.offset(offset).limit(limit).all()


@router.post("")
def create_opportunity(
    payload: schemas.OpportunityCreate,
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    opp = models.Opportunity(
        opp_code=_gen_opp_code(),
        title=payload.title,
        agency_name=payload.agency_name,
        agency_type=payload.agency_type,
        description=payload.description,
        naics_code=payload.naics_code,
        psc_code=payload.psc_code,
        set_aside=payload.set_aside,
        estimated_value_low=payload.estimated_value_low,
        estimated_value_high=payload.estimated_value_high,
        location_city=payload.location_city,
        location_state=payload.location_state,
        posted_date=payload.posted_date,
        due_date=payload.due_date,
        source_type=payload.source_type or "manual",
        source_url=payload.source_url,
        solicitation_number=payload.solicitation_number,
        contract_type=payload.contract_type,
        status="new",
    )
    db.add(opp)
    db.commit()
    db.refresh(opp)
    return opp


@router.get("/{opp_id}")
def get_opportunity(opp_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    return opp


@router.delete("/{opp_id}")
def delete_opportunity(opp_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    for att in opp.attachments:
        try:
            if att.stored_path and os.path.exists(att.stored_path):
                os.remove(att.stored_path)
        except Exception:
            pass
    db.delete(opp)
    db.commit()
    return {"ok": True, "deleted": opp_id}


# ─── File Upload ──────────────────────────────────────────

@router.post("/{opp_id}/attachments")
async def upload_attachment(
    opp_id: int, files: List[UploadFile] = File(...), category: str = Form("solicitation"),
    user: str = Depends(require_auth), db: Session = Depends(get_db),
):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")

    upload_dir = os.path.join(UPLOAD_DIR, str(opp_id))
    os.makedirs(upload_dir, exist_ok=True)

    results = []
    for file in files:
        ext = (file.filename or "").rsplit(".", 1)[-1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            continue
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            continue
        safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
        filepath = os.path.join(upload_dir, safe_name)
        with open(filepath, "wb") as f:
            f.write(content)
        extracted = _extract_text(filepath, ext)
        att = models.OpportunityAttachment(
            opportunity_id=opp_id, filename=file.filename, stored_path=filepath,
            file_type=ext, file_size=len(content), category=category,
            extracted_text=extracted if extracted.strip() else None,
            original_filename=file.filename,
        )
        db.add(att)
        db.commit()
        db.refresh(att)
        results.append({"id": att.id, "filename": att.filename, "file_type": att.file_type,
                        "file_size": att.file_size, "has_extracted_text": bool(att.extracted_text)})

    return {"files": results, "count": len(results)}


@router.get("/{opp_id}/attachments")
def list_attachments(opp_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    return [{"id": a.id, "filename": a.filename, "file_type": a.file_type, "file_size": a.file_size,
             "category": a.category, "has_extracted_text": bool(a.extracted_text), "uploaded_at": a.uploaded_at}
            for a in opp.attachments]


@router.delete("/{opp_id}/attachments/{att_id}")
def delete_attachment(opp_id: int, att_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    att = db.query(models.OpportunityAttachment).filter(
        models.OpportunityAttachment.id == att_id, models.OpportunityAttachment.opportunity_id == opp_id).first()
    if not att:
        raise HTTPException(404, "Attachment not found")
    try:
        if att.stored_path and os.path.exists(att.stored_path): os.remove(att.stored_path)
    except Exception: pass
    db.delete(att)
    db.commit()
    return {"ok": True}


# ─── AI Analysis ──────────────────────────────────────────

@router.post("/{opp_id}/analyze")
def analyze_opportunity(opp_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    attachments_text = _get_attachments_text(opp)
    context = _build_opp_context(opp, attachments_text)

    # Add business profile context if available
    profile_text = _get_business_profile(db, user)
    if profile_text:
        context += f"\n\n--- Your Business Profile ---\n{profile_text}"

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise HTTPException(400, "OpenAI API key not configured. Set OPENAI_API_KEY in .env")

    results = {}
    summary = _ai_executive_summary(context)
    if summary: opp.ai_summary = summary; results["summary"] = json.loads(summary)
    reqs = _ai_requirements(context)
    if reqs: opp.ai_requirements = reqs; results["requirements"] = json.loads(reqs)
    evalf = _ai_evaluation_factors(context)
    if evalf: opp.ai_evaluation_factors = evalf; results["evaluation_factors"] = json.loads(evalf)
    risks = _ai_risk_flags(context)
    if risks: opp.ai_risk_flags = risks; results["risk_flags"] = json.loads(risks)
    compliance = _ai_compliance_checklist(context)
    if compliance: opp.ai_compliance_checklist = compliance; results["compliance"] = json.loads(compliance)
    strategy = _ai_bid_strategy(context)
    if strategy: opp.ai_bid_strategy = strategy; results["strategy"] = json.loads(strategy)

    prior = json.dumps(results, indent=2)[:4000]
    rec = _ai_bid_recommendation(context, prior)
    if rec:
        try:
            rec_data = json.loads(rec)
            opp.ai_bid_recommendation = rec_data.get("recommendation", "conditional")
            opp.ai_confidence_score = rec_data.get("confidence", 0.5)
            results["recommendation"] = rec_data
        except Exception:
            results["recommendation"] = rec

    opp.ai_analyzed_at = datetime.now(timezone.utc)
    opp.status = "analyzed"
    db.commit()
    db.refresh(opp)
    return {"ok": True, "opp_id": opp.id, "status": opp.status,
            "ai_bid_recommendation": opp.ai_bid_recommendation,
            "ai_confidence_score": opp.ai_confidence_score, "analysis": results}


# ─── Ask AI ────────────────────────────────────────────────

@router.post("/{opp_id}/ask")
def ask_about_opportunity(opp_id: int, body: dict = Body(...), user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    question = body.get("question", "").strip()
    if not question:
        raise HTTPException(400, "Question is required")
    attachments_text = _get_attachments_text(opp)
    context = _build_opp_context(opp, attachments_text)
    prior = ""
    if opp.ai_summary: prior += f"\nSummary: {opp.ai_summary[:500]}"
    if opp.ai_bid_recommendation: prior += f"\nRecommendation: {opp.ai_bid_recommendation}"

    profile_text = _get_business_profile(db, user)
    profile_section = f"\n\n--- Your Business Profile ---\n{profile_text}" if profile_text else ""

    system = f"""You are a senior government contracting advisor for SentriBiD.
Answer questions about this opportunity. Be specific and actionable.
{context}\n{prior}{profile_section}"""
    answer = _call_openai(system, question, json_mode=False)
    if not answer:
        return {"answer": "AI is not available. Please configure your OpenAI API key."}
    return {"answer": answer}


# ─── Convert to Bid (Smart - with AI pricing) ─────────────

@router.post("/{opp_id}/convert")
def convert_to_bid(opp_id: int, user: str = Depends(require_auth), db: Session = Depends(get_db)):
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")
    if opp.converted_bid_id:
        raise HTTPException(400, f"Already converted to bid ID {opp.converted_bid_id}")

    # Use AI to suggest line items and pricing
    ai_items = []
    ai_labor = []
    ai_overhead = {}
    context = _build_opp_context(opp, _get_attachments_text(opp))

    profile_text = _get_business_profile(db, user)
    if profile_text:
        context += f"\n\n--- Your Business Profile ---\n{profile_text}"

    pricing_prompt = """Based on this government solicitation and the business profile, estimate a detailed cost breakdown.
Return JSON: {
  "items": [{"name": "...", "description": "...", "quantity": 1, "unit_cost": 0.0}],
  "labor": [{"labor_type": "...", "hourly_rate": 0.0, "hours": 0.0, "workers": 1}],
  "overhead": {"insurance_allocation": 0, "admin_time_cost": 0, "bonding_compliance_cost": 0, "misc_overhead": 0},
  "total_estimated_cost": 0,
  "suggested_bid_price": 0,
  "profit_margin_pct": 0,
  "pricing_notes": "..."
}
Be realistic with government contract pricing. Include all relevant cost categories."""

    pricing_raw = _call_openai(
        "You are a government contract pricing specialist. Provide realistic cost estimates.",
        f"{context}\n\n{pricing_prompt}"
    )

    if pricing_raw:
        try:
            pricing = json.loads(pricing_raw)
            ai_items = pricing.get("items", [])
            ai_labor = pricing.get("labor", [])
            ai_overhead = pricing.get("overhead", {})
        except Exception:
            pass

    # Build notes from AI analysis
    notes = f"Converted from opportunity {opp.opp_code}\n"
    if opp.ai_summary:
        try:
            s = json.loads(opp.ai_summary)
            notes += f"\nSummary: {s.get('summary', '')[:500]}"
        except Exception: pass
    if opp.ai_bid_recommendation:
        notes += f"\nAI Recommendation: {opp.ai_bid_recommendation}"
    if pricing_raw:
        try:
            p = json.loads(pricing_raw)
            if p.get("pricing_notes"):
                notes += f"\nPricing Notes: {p['pricing_notes'][:300]}"
        except Exception: pass

    risk_level = 3
    if opp.ai_risk_flags:
        try:
            r = json.loads(opp.ai_risk_flags)
            level = r.get("overall_risk_level", "medium")
            risk_level = {"low": 1, "medium": 3, "high": 5}.get(level, 3)
        except Exception: pass

    bid_code = f"SB-{datetime.now().year}-{uuid.uuid4().hex[:8].upper()}"
    bid = models.Bid(
        bid_code=bid_code, contract_title=opp.title, agency_name=opp.agency_name,
        agency_type=opp.agency_type or "federal", solicitation_number=opp.solicitation_number,
        contract_type=opp.contract_type or "service", delivery_distance_miles=0,
        deadline_date=opp.due_date.date() if opp.due_date else date.today(),
        urgency_level=3, competition_level="medium", risk_level=risk_level,
        desired_profit_mode="balanced", status="draft", notes=notes[:2000],
    )
    db.add(bid)
    db.flush()

    # Add AI-suggested line items
    for item in ai_items[:10]:
        bi = models.BidItem(
            bid_id=bid.id, name=item.get("name", "Item"),
            description=item.get("description", ""),
            quantity=float(item.get("quantity", 1)),
            unit_cost=float(item.get("unit_cost", 0)),
        )
        db.add(bi)

    # Add AI-suggested labor
    for lab in ai_labor[:5]:
        bl = models.BidLaborLine(
            bid_id=bid.id, labor_type=lab.get("labor_type", "other"),
            hourly_rate=float(lab.get("hourly_rate", 0)),
            hours=float(lab.get("hours", 0)),
            workers=int(lab.get("workers", 1)),
        )
        db.add(bl)

    # Add AI-suggested overhead
    if ai_overhead:
        bo = models.BidOverhead(
            bid_id=bid.id,
            insurance_allocation=float(ai_overhead.get("insurance_allocation", 0)),
            admin_time_cost=float(ai_overhead.get("admin_time_cost", 0)),
            bonding_compliance_cost=float(ai_overhead.get("bonding_compliance_cost", 0)),
            misc_overhead=float(ai_overhead.get("misc_overhead", 0)),
        )
        db.add(bo)

    opp.converted_bid_id = bid.id
    opp.status = "converted"
    db.commit()
    db.refresh(bid)

    return {"ok": True, "bid_id": bid.id, "bid_code": bid.bid_code, "opp_id": opp.id,
            "items_added": len(ai_items), "labor_added": len(ai_labor),
            "message": f"Opportunity converted to bid {bid.bid_code} with AI-estimated pricing"}


# ─── Generate Proposal (PDF/DOCX) ─────────────────────────

@router.post("/{opp_id}/generate-proposal")
def generate_proposal(
    opp_id: int,
    body: dict = Body(default={}),
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Generate a professional proposal document based on opportunity analysis + business profile."""
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")

    format_type = body.get("format", "pdf")  # pdf or docx
    attachments_text = _get_attachments_text(opp)
    context = _build_opp_context(opp, attachments_text)

    profile_text = _get_business_profile(db, user)
    if not profile_text:
        raise HTTPException(400, "Please complete your Business Profile before generating a proposal.")

    # Get the user's profile object for the generator
    u = db.query(models.User).filter(models.User.email == user).first()
    company_name = u.profile.company_name if (u and u.profile and u.profile.company_name) else "Your Company"
    profile_obj = u.profile if u else None

    # Get bid if converted
    bid_obj = None
    if opp.converted_bid_id:
        bid_obj = db.query(models.Bid).filter(models.Bid.id == opp.converted_bid_id).first()

    # Build pricing context for the AI
    pricing_context = ""
    if bid_obj:
        item_total = sum(getattr(i, 'quantity', 0) * getattr(i, 'unit_cost', 0) for i in (bid_obj.items or []))
        labor_total = sum(getattr(l, 'hourly_rate', 0) * getattr(l, 'hours', 0) * getattr(l, 'workers', 1) for l in (bid_obj.labor_lines or []))
        pricing_context = f"\nPricing: Items=${item_total:,.2f}, Labor=${labor_total:,.2f}, Total=${item_total + labor_total:,.2f}"

    # Gather all prior analysis
    analysis_context = ""
    if opp.ai_summary: analysis_context += f"\nSummary: {opp.ai_summary[:800]}"
    if opp.ai_requirements: analysis_context += f"\nRequirements: {opp.ai_requirements[:800]}"
    if opp.ai_evaluation_factors: analysis_context += f"\nEvaluation: {opp.ai_evaluation_factors[:500]}"
    if opp.ai_bid_strategy: analysis_context += f"\nStrategy: {opp.ai_bid_strategy[:500]}"
    if opp.ai_compliance_checklist: analysis_context += f"\nCompliance: {opp.ai_compliance_checklist[:500]}"

    # Ask AI to write the full proposal — Claude first (best quality), Gemini fallback, OpenAI last
    proposal_text = None

    # 1. Try Claude first (best reasoning and quality)
    try:
        from ..claude_ai import call_claude
        proposal_prompt_claude = f"""Write a complete, professional government contract proposal response.

OPPORTUNITY:
{context}

PRIOR ANALYSIS:
{analysis_context}

YOUR COMPANY PROFILE:
{profile_text}
{pricing_context}

Write the proposal with these sections. Each section should be 3-5 detailed paragraphs:
1. COVER LETTER - Professional letter to the contracting officer
2. EXECUTIVE SUMMARY - Why {company_name} is the best choice
3. TECHNICAL APPROACH - How we'll deliver the work (detailed, specific to THIS solicitation)
4. MANAGEMENT PLAN - Team structure, key personnel, communication plan
5. PAST PERFORMANCE - Relevant contracts completed
6. STAFFING PLAN - Who will work on this, their qualifications
7. QUALITY ASSURANCE - How we ensure quality delivery
8. PRICING NARRATIVE - Justify the pricing approach

Make it compelling, specific, and reference the company's actual capabilities.
Use the company name "{company_name}" throughout."""

        proposal_text = call_claude(
            prompt=proposal_prompt_claude,
            system_instruction=f"You are an expert government proposal writer for {company_name}. Write winning, detailed proposals in proper government contracting language.",
            max_tokens=8000,
            temperature=0.3,
        )
        if proposal_text:
            logger.info("Proposal generated with Claude AI")
    except Exception as e:
        logger.warning(f"Claude proposal failed: {e}")

    # 2. Try Gemini (good for long docs)
    if not proposal_text:
        try:
            from ..gemini_ai import call_gemini_proposal
            proposal_text = call_gemini_proposal(
                opportunity_context=context,
                analysis_context=analysis_context,
                profile_text=profile_text,
                pricing_context=pricing_context,
                company_name=company_name,
            )
            if proposal_text:
                logger.info("Proposal generated with Gemini AI")
        except Exception as e:
            logger.warning(f"Gemini proposal failed: {e}")

    # 3. Fallback to OpenAI
    if not proposal_text:
        proposal_prompt = f"""Write a complete, professional government contract proposal response.

OPPORTUNITY:
{context}

PRIOR ANALYSIS:
{analysis_context}

YOUR COMPANY PROFILE:
{profile_text}
{pricing_context}

Write the proposal with these sections:
1. COVER LETTER - Professional letter to the contracting officer
2. EXECUTIVE SUMMARY - Why we're the best choice (2-3 paragraphs)
3. TECHNICAL APPROACH - How we'll deliver the work (detailed, specific to THIS solicitation)
4. MANAGEMENT PLAN - Team structure, key personnel, communication plan
5. PAST PERFORMANCE - Relevant contracts we've completed (use actual data from the business profile)
6. STAFFING PLAN - Who will work on this, their qualifications
7. QUALITY ASSURANCE - How we ensure quality delivery
8. PRICING NARRATIVE - Justify the pricing approach, reference actual costs if available

Make it compelling, specific to THIS opportunity, and reference the company's actual capabilities.
Use professional government contracting language. Be detailed and thorough.
Each section should be 2-4 paragraphs. Use the company's real name "{company_name}" throughout."""

        proposal_text = _call_openai(
            "You are an expert government proposal writer who wins contracts. Write compelling, detailed proposals.",
            proposal_prompt,
            json_mode=False,
        )

    if not proposal_text:
        raise HTTPException(500, "AI proposal generation failed. Check your CLAUDE_API_KEY, GEMINI_API_KEY, or OPENAI_API_KEY in .env")

    from ..proposal_generator import generate_pdf_proposal, generate_docx_proposal

    if format_type == "docx":
        filepath = generate_docx_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid_obj)
    else:
        filepath = generate_pdf_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid_obj)

    media = "application/pdf" if format_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ext = format_type
    return FileResponse(filepath, media_type=media,
                       filename=f"Proposal-{opp.opp_code}.{ext}",
                       headers={"Content-Disposition": f"attachment; filename=Proposal-{opp.opp_code}.{ext}"})



# ─── Upload & Analyze (Quick Flow) ───────────────────────

@router.post("/upload-and-analyze")
async def upload_and_analyze(
    files: List[UploadFile] = File(...), title: str = Form(""), agency_name: str = Form("Unknown Agency"),
    source_url: str = Form(""), user: str = Depends(require_auth), db: Session = Depends(get_db),
):
    if not files:
        raise HTTPException(400, "No files provided")

    # Use first file's name as default title
    first_file = files[0]
    opp = models.Opportunity(
        opp_code=_gen_opp_code(),
        title=title or (first_file.filename or "Uploaded Solicitation").rsplit(".", 1)[0],
        agency_name=agency_name, source_type="upload", source_url=source_url or None, status="new",
    )
    db.add(opp)
    db.flush()

    upload_dir = os.path.join(UPLOAD_DIR, str(opp.id))
    os.makedirs(upload_dir, exist_ok=True)

    all_text = []
    file_count = 0

    for file in files:
        ext = (file.filename or "").rsplit(".", 1)[-1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            continue
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            continue

        safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
        filepath = os.path.join(upload_dir, safe_name)
        with open(filepath, "wb") as f:
            f.write(content)

        extracted = _extract_text(filepath, ext)
        att = models.OpportunityAttachment(
            opportunity_id=opp.id, filename=file.filename, stored_path=filepath,
            file_type=ext, file_size=len(content), category="solicitation",
            extracted_text=extracted if extracted.strip() else None,
            original_filename=file.filename,
        )
        db.add(att)
        file_count += 1
        if extracted.strip():
            all_text.append(f"[{file.filename}]\n{extracted}")

    db.commit()
    db.refresh(opp)

    return {
        "ok": True, "opp_id": opp.id, "opp_code": opp.opp_code, "title": opp.title,
        "files_uploaded": file_count,
        "has_text": bool(all_text),
        "message": f"{file_count} file(s) uploaded. Click Analyze to run AI analysis.",
    }


# ─── BID AUTOPILOT v0.7.0 ────────────────────────────────
# One-click: Upload RFP → Shred → Compliance → War Room → Bid → Proposal

@router.post("/autopilot-upload")
async def autopilot_upload(
    files: List[UploadFile] = File(...),
    title: str = Form(""),
    agency_name: str = Form("Unknown Agency"),
    format_type: str = Form("pdf"),
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """
    BID AUTOPILOT — Upload RFP file(s) and get a complete submission package in one call.
    Accepts multiple files (PDF, DOCX, XLSX, CSV, TXT).

    Pipeline: Upload → Extract → AI Analysis → Shred RFP → Compliance Matrix →
              War Room → Convert to Bid → Generate Proposal (PDF/DOCX)

    Returns all analysis results + a downloadable proposal file.
    """
    if not files:
        raise HTTPException(400, "No files provided")

    steps_completed = []
    errors = []

    # ── Step 1: Upload & Extract ALL files ────────────────
    first_file = files[0]
    opp = models.Opportunity(
        opp_code=_gen_opp_code(),
        title=title or (first_file.filename or "Uploaded RFP").rsplit(".", 1)[0],
        agency_name=agency_name,
        source_type="upload",
        status="new",
    )
    db.add(opp)
    db.flush()

    upload_dir = os.path.join(UPLOAD_DIR, str(opp.id))
    os.makedirs(upload_dir, exist_ok=True)

    all_extracted = []
    file_count = 0

    for file in files:
        ext = (file.filename or "").rsplit(".", 1)[-1].lower()
        if ext not in ALLOWED_EXTENSIONS:
            errors.append(f"Skipped {file.filename} — unsupported type .{ext}")
            continue
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            errors.append(f"Skipped {file.filename} — exceeds 20MB limit")
            continue

        safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
        filepath = os.path.join(upload_dir, safe_name)
        with open(filepath, "wb") as f_out:
            f_out.write(content)

        extracted = _extract_text(filepath, ext)
        att = models.OpportunityAttachment(
            opportunity_id=opp.id, filename=file.filename, stored_path=filepath,
            file_type=ext, file_size=len(content), category="solicitation",
            extracted_text=extracted if extracted.strip() else None,
            original_filename=file.filename,
        )
        db.add(att)
        file_count += 1
        if extracted.strip():
            all_extracted.append(f"[{file.filename}]\n{extracted}")

    db.commit()
    db.refresh(opp)
    steps_completed.append("upload")

    combined_text = "\n\n".join(all_extracted)
    if not combined_text.strip():
        raise HTTPException(400, f"Could not extract text from {file_count} file(s). Try text-based PDFs or DOCX files.")

    # Get profile context
    profile_text = _get_business_profile(db, user)
    if not profile_text:
        raise HTTPException(400, "Complete your Business Profile first — Autopilot needs it to generate proposals.")

    u = db.query(models.User).filter(models.User.email == user).first()
    company_name = u.profile.company_name if (u and u.profile and u.profile.company_name) else "Your Company"
    profile_obj = u.profile if u else None

    attachments_text = combined_text[:25000]
    context = _build_opp_context(opp, attachments_text)
    if profile_text:
        context += f"\n\n--- Your Business Profile ---\n{profile_text}"

    # ── Step 2: AI Analysis (6 sub-analyses) ──────────────
    analysis_results = {}
    try:
        summary = _ai_executive_summary(context)
        if summary:
            opp.ai_summary = summary
            analysis_results["summary"] = json.loads(summary)

        reqs = _ai_requirements(context)
        if reqs:
            opp.ai_requirements = reqs
            analysis_results["requirements"] = json.loads(reqs)

        evalf = _ai_evaluation_factors(context)
        if evalf:
            opp.ai_evaluation_factors = evalf
            analysis_results["evaluation_factors"] = json.loads(evalf)

        risks = _ai_risk_flags(context)
        if risks:
            opp.ai_risk_flags = risks
            analysis_results["risk_flags"] = json.loads(risks)

        compliance = _ai_compliance_checklist(context)
        if compliance:
            opp.ai_compliance_checklist = compliance
            analysis_results["compliance"] = json.loads(compliance)

        strategy = _ai_bid_strategy(context)
        if strategy:
            opp.ai_bid_strategy = strategy
            analysis_results["strategy"] = json.loads(strategy)

        prior = json.dumps(analysis_results, indent=2)[:4000]
        rec = _ai_bid_recommendation(context, prior)
        if rec:
            try:
                rec_data = json.loads(rec)
                opp.ai_bid_recommendation = rec_data.get("recommendation", "conditional")
                opp.ai_confidence_score = rec_data.get("confidence", 0.5)
                analysis_results["recommendation"] = rec_data
            except Exception:
                pass

        opp.ai_analyzed_at = datetime.now(timezone.utc)
        opp.status = "analyzed"
        db.commit()
        steps_completed.append("analysis")
    except Exception as e:
        errors.append(f"Analysis: {str(e)}")
        logger.warning(f"Autopilot analysis error: {e}")

    # ── Step 3: Shred RFP ─────────────────────────────────
    shredded = {}
    try:
        from ..compliance_engine import shred_rfp
        shredded = shred_rfp(combined_text, opportunity_title=opp.title)
        if not shredded.get("error"):
            if hasattr(opp, 'shredded_rfp'):
                opp.shredded_rfp = json.dumps(shredded)
                db.commit()
            steps_completed.append("shred")
        else:
            errors.append(f"Shred: {shredded.get('error')}")
    except Exception as e:
        errors.append(f"Shred: {str(e)}")
        logger.warning(f"Autopilot shred error: {e}")

    # ── Step 4: Compliance Matrix ─────────────────────────
    compliance_matrix = {}
    try:
        if shredded and shredded.get("requirements"):
            from ..compliance_engine import generate_compliance_matrix
            compliance_matrix = generate_compliance_matrix(shredded, profile_text or "No profile")
            if not compliance_matrix.get("error"):
                if hasattr(opp, 'compliance_matrix'):
                    opp.compliance_matrix = json.dumps(compliance_matrix)
                    db.commit()
                steps_completed.append("compliance")
            else:
                errors.append(f"Compliance: {compliance_matrix.get('error')}")
        else:
            errors.append("Compliance: Skipped (no shredded requirements)")
    except Exception as e:
        errors.append(f"Compliance: {str(e)}")
        logger.warning(f"Autopilot compliance error: {e}")

    # ── Step 5: War Room ──────────────────────────────────
    war_room = {}
    try:
        from ..win_predictor import run_war_room
        from ..sam_connector import search_awards

        opp_context = f"Title: {opp.title}\nAgency: {opp.agency_name}\nNAICS: {opp.naics_code}\nDue: {opp.due_date}\n"
        opp_context += f"\nDocument:\n{combined_text[:3000]}"

        historical = []
        search_term = opp.naics_code or (opp.title.split()[0] if opp.title else "")
        if search_term:
            try:
                award_result = search_awards(keyword=search_term, naics_code=opp.naics_code or "", limit=10)
                historical = award_result.get("awards", [])
            except Exception:
                pass

        war_room = run_war_room(opp_context, profile_text or "No profile", historical, shredded or None)
        if not war_room.get("error"):
            if hasattr(opp, 'war_room_analysis'):
                opp.war_room_analysis = json.dumps(war_room)
                db.commit()
            steps_completed.append("war_room")
        else:
            errors.append(f"War Room: {war_room.get('error')}")
    except Exception as e:
        errors.append(f"War Room: {str(e)}")
        logger.warning(f"Autopilot war room error: {e}")

    # ── Step 6: Convert to Bid ────────────────────────────
    bid = None
    try:
        ai_items = []
        ai_labor = []
        ai_overhead_data = {}

        pricing_raw = _call_openai(
            "You are a government contract pricing specialist. Provide realistic cost estimates.",
            f"""{context}\n\nEstimate a detailed cost breakdown. Return JSON: {{
  "items": [{{"name": "...", "description": "...", "quantity": 1, "unit_cost": 0.0}}],
  "labor": [{{"labor_type": "...", "hourly_rate": 0.0, "hours": 0.0, "workers": 1}}],
  "overhead": {{"insurance_allocation": 0, "admin_time_cost": 0, "bonding_compliance_cost": 0, "misc_overhead": 0}},
  "total_estimated_cost": 0, "suggested_bid_price": 0}}"""
        )
        if pricing_raw:
            try:
                pricing = json.loads(pricing_raw)
                ai_items = pricing.get("items", [])
                ai_labor = pricing.get("labor", [])
                ai_overhead_data = pricing.get("overhead", {})
            except Exception:
                pass

        risk_level = 3
        if opp.ai_risk_flags:
            try:
                r = json.loads(opp.ai_risk_flags)
                level = r.get("overall_risk_level", "medium")
                risk_level = {"low": 1, "medium": 3, "high": 5}.get(level, 3)
            except Exception:
                pass

        bid_code = f"SB-{datetime.now().year}-{uuid.uuid4().hex[:8].upper()}"
        bid = models.Bid(
            bid_code=bid_code, contract_title=opp.title, agency_name=opp.agency_name,
            agency_type=opp.agency_type or "federal", solicitation_number=opp.solicitation_number,
            contract_type=opp.contract_type or "service", delivery_distance_miles=0,
            deadline_date=opp.due_date.date() if opp.due_date else date.today(),
            urgency_level=3, competition_level="medium", risk_level=risk_level,
            desired_profit_mode="balanced", status="draft",
            notes=f"Auto-generated by Bid Autopilot from {opp.opp_code}",
        )
        db.add(bid)
        db.flush()

        for item in ai_items[:10]:
            db.add(models.BidItem(
                bid_id=bid.id, name=item.get("name", "Item"),
                description=item.get("description", ""),
                quantity=float(item.get("quantity", 1)),
                unit_cost=float(item.get("unit_cost", 0)),
            ))
        for lab in ai_labor[:5]:
            db.add(models.BidLaborLine(
                bid_id=bid.id, labor_type=lab.get("labor_type", "other"),
                hourly_rate=float(lab.get("hourly_rate", 0)),
                hours=float(lab.get("hours", 0)),
                workers=int(lab.get("workers", 1)),
            ))
        if ai_overhead_data:
            db.add(models.BidOverhead(
                bid_id=bid.id,
                insurance_allocation=float(ai_overhead_data.get("insurance_allocation", 0)),
                admin_time_cost=float(ai_overhead_data.get("admin_time_cost", 0)),
                bonding_compliance_cost=float(ai_overhead_data.get("bonding_compliance_cost", 0)),
                misc_overhead=float(ai_overhead_data.get("misc_overhead", 0)),
            ))

        opp.converted_bid_id = bid.id
        opp.status = "converted"
        db.commit()
        db.refresh(bid)
        steps_completed.append("bid")
    except Exception as e:
        errors.append(f"Bid: {str(e)}")
        logger.warning(f"Autopilot bid error: {e}")

    # ── Step 7: Generate Proposal ─────────────────────────
    proposal_path = None
    try:
        analysis_context = ""
        if opp.ai_summary:
            analysis_context += f"\nSummary: {opp.ai_summary[:800]}"
        if opp.ai_requirements:
            analysis_context += f"\nRequirements: {opp.ai_requirements[:800]}"
        if opp.ai_evaluation_factors:
            analysis_context += f"\nEvaluation: {opp.ai_evaluation_factors[:500]}"
        if opp.ai_bid_strategy:
            analysis_context += f"\nStrategy: {opp.ai_bid_strategy[:500]}"

        pricing_context = ""
        if bid:
            item_total = sum(getattr(i, 'quantity', 0) * getattr(i, 'unit_cost', 0) for i in (bid.items or []))
            labor_total = sum(getattr(l, 'hourly_rate', 0) * getattr(l, 'hours', 0) * getattr(l, 'workers', 1) for l in (bid.labor_lines or []))
            pricing_context = f"\nPricing: Items=${item_total:,.2f}, Labor=${labor_total:,.2f}, Total=${item_total + labor_total:,.2f}"

        # Claude-first proposal generation
        proposal_text = None
        try:
            from ..claude_ai import call_claude
            proposal_text = call_claude(
                prompt=f"""Write a complete, professional government contract proposal response.

OPPORTUNITY:
{context[:4000]}

PRIOR ANALYSIS:
{analysis_context}

WAR ROOM INSIGHTS:
{json.dumps(war_room.get('our_win_strategy', {}), default=str)[:1500] if war_room else 'N/A'}

COMPLIANCE GAPS:
{json.dumps(compliance_matrix.get('critical_gaps', []), default=str)[:800] if compliance_matrix else 'N/A'}

YOUR COMPANY PROFILE:
{profile_text}
{pricing_context}

Write these sections (3-5 paragraphs each):
1. COVER LETTER 2. EXECUTIVE SUMMARY 3. TECHNICAL APPROACH
4. MANAGEMENT PLAN 5. PAST PERFORMANCE 6. STAFFING PLAN
7. QUALITY ASSURANCE 8. PRICING NARRATIVE

Use company name "{company_name}" throughout. Make it compelling and specific.""",
                system_instruction=f"You are an expert government proposal writer for {company_name}. Write winning proposals.",
                max_tokens=8000,
                temperature=0.3,
            )
        except Exception:
            pass

        if not proposal_text:
            try:
                from ..gemini_ai import call_gemini_proposal
                proposal_text = call_gemini_proposal(
                    opportunity_context=context, analysis_context=analysis_context,
                    profile_text=profile_text or "", pricing_context=pricing_context,
                    company_name=company_name,
                )
            except Exception:
                pass

        if not proposal_text:
            proposal_text = _call_openai(
                f"You are an expert proposal writer for {company_name}.",
                f"Write a complete government proposal for:\n{context[:3000]}\n\nProfile:\n{profile_text[:2000]}\n\nInclude: Cover Letter, Executive Summary, Technical Approach, Management Plan, Past Performance, Staffing, QA, Pricing Narrative.",
                json_mode=False,
            )

        if proposal_text:
            from ..proposal_generator import generate_pdf_proposal, generate_docx_proposal
            if format_type == "docx":
                proposal_path = generate_docx_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid)
            else:
                proposal_path = generate_pdf_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid)
            steps_completed.append("proposal")
        else:
            errors.append("Proposal: All AI engines failed to generate proposal text")
    except Exception as e:
        errors.append(f"Proposal: {str(e)}")
        logger.warning(f"Autopilot proposal error: {e}")

    # ── Return Results ────────────────────────────────────
    result = {
        "ok": True,
        "opp_id": opp.id,
        "opp_code": opp.opp_code,
        "bid_id": bid.id if bid else None,
        "bid_code": bid.bid_code if bid else None,
        "steps_completed": steps_completed,
        "steps_total": 7,
        "errors": errors,
        "analysis": {
            "recommendation": opp.ai_bid_recommendation,
            "confidence": opp.ai_confidence_score,
            "summary": analysis_results.get("summary"),
        },
        "war_room": {
            "win_probability": war_room.get("win_probability", {}).get("score", 0) if war_room else 0,
            "executive_brief": war_room.get("executive_brief", "") if war_room else "",
            "bottom_line": war_room.get("bottom_line", "") if war_room else "",
        },
        "compliance": {
            "score": compliance_matrix.get("compliance_score", 0) if compliance_matrix else 0,
            "gaps": len(compliance_matrix.get("critical_gaps", [])) if compliance_matrix else 0,
        },
        "has_proposal": proposal_path is not None,
        "proposal_download": f"/opportunities/{opp.id}/autopilot-download?format={format_type}" if proposal_path else None,
    }

    # Store the proposal path for download
    if proposal_path:
        # Store in a simple way — save path to opp notes or a temp location
        if not opp.capture_notes:
            opp.capture_notes = ""
        opp.capture_notes = f"autopilot_proposal_path={proposal_path}\n" + (opp.capture_notes or "")
        db.commit()

    return result


@router.post("/{opp_id}/autopilot")
def autopilot_existing(
    opp_id: int,
    body: dict = Body(default={}),
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Run Bid Autopilot on an existing opportunity (must have attachments with extracted text)."""
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")

    attachments_text = _get_attachments_text(opp)
    if not attachments_text.strip() and not opp.description:
        raise HTTPException(400, "This opportunity has no documents or description. Upload an RFP first.")

    profile_text = _get_business_profile(db, user)
    if not profile_text:
        raise HTTPException(400, "Complete your Business Profile first.")

    format_type = body.get("format", "pdf")
    u = db.query(models.User).filter(models.User.email == user).first()
    company_name = u.profile.company_name if (u and u.profile and u.profile.company_name) else "Your Company"

    # Run same pipeline as autopilot_upload but skip the upload step
    # For brevity, call the individual endpoints internally
    steps_completed = []
    errors = []
    context = _build_opp_context(opp, attachments_text)
    if profile_text:
        context += f"\n\n--- Your Business Profile ---\n{profile_text}"

    # Analysis
    if not opp.ai_summary:
        try:
            summary = _ai_executive_summary(context)
            if summary:
                opp.ai_summary = summary
            reqs = _ai_requirements(context)
            if reqs:
                opp.ai_requirements = reqs
            evalf = _ai_evaluation_factors(context)
            if evalf:
                opp.ai_evaluation_factors = evalf
            risks = _ai_risk_flags(context)
            if risks:
                opp.ai_risk_flags = risks
            compliance_raw = _ai_compliance_checklist(context)
            if compliance_raw:
                opp.ai_compliance_checklist = compliance_raw
            strategy = _ai_bid_strategy(context)
            if strategy:
                opp.ai_bid_strategy = strategy

            prior = json.dumps({"summary": opp.ai_summary, "strategy": opp.ai_bid_strategy}, default=str)[:3000]
            rec = _ai_bid_recommendation(context, prior)
            if rec:
                try:
                    rec_data = json.loads(rec)
                    opp.ai_bid_recommendation = rec_data.get("recommendation", "conditional")
                    opp.ai_confidence_score = rec_data.get("confidence", 0.5)
                except Exception:
                    pass

            opp.ai_analyzed_at = datetime.now(timezone.utc)
            opp.status = "analyzed"
            db.commit()
            steps_completed.append("analysis")
        except Exception as e:
            errors.append(f"Analysis: {str(e)}")
    else:
        steps_completed.append("analysis")

    # Shred
    doc_text = attachments_text or opp.description or ""
    shredded = {}
    if not getattr(opp, 'shredded_rfp', None):
        try:
            from ..compliance_engine import shred_rfp
            shredded = shred_rfp(doc_text, opportunity_title=opp.title)
            if not shredded.get("error"):
                opp.shredded_rfp = json.dumps(shredded)
                db.commit()
                steps_completed.append("shred")
        except Exception as e:
            errors.append(f"Shred: {str(e)}")
    else:
        shredded = json.loads(opp.shredded_rfp)
        steps_completed.append("shred")

    # Compliance
    compliance_matrix = {}
    if not getattr(opp, 'compliance_matrix', None) and shredded.get("requirements"):
        try:
            from ..compliance_engine import generate_compliance_matrix
            compliance_matrix = generate_compliance_matrix(shredded, profile_text)
            if not compliance_matrix.get("error"):
                opp.compliance_matrix = json.dumps(compliance_matrix)
                db.commit()
                steps_completed.append("compliance")
        except Exception as e:
            errors.append(f"Compliance: {str(e)}")
    elif getattr(opp, 'compliance_matrix', None):
        compliance_matrix = json.loads(opp.compliance_matrix)
        steps_completed.append("compliance")

    # War Room
    war_room = {}
    if not getattr(opp, 'war_room_analysis', None):
        try:
            from ..win_predictor import run_war_room
            opp_context = f"Title: {opp.title}\nAgency: {opp.agency_name}\nNAICS: {opp.naics_code}\n"
            opp_context += f"\nDocument:\n{doc_text[:3000]}"
            war_room = run_war_room(opp_context, profile_text, [], shredded or None)
            if not war_room.get("error"):
                opp.war_room_analysis = json.dumps(war_room)
                db.commit()
                steps_completed.append("war_room")
        except Exception as e:
            errors.append(f"War Room: {str(e)}")
    else:
        war_room = json.loads(opp.war_room_analysis)
        steps_completed.append("war_room")

    # Convert to Bid (skip if already converted)
    bid = None
    if opp.converted_bid_id:
        bid = db.query(models.Bid).filter(models.Bid.id == opp.converted_bid_id).first()
        steps_completed.append("bid")
    else:
        try:
            bid_code = f"SB-{datetime.now().year}-{uuid.uuid4().hex[:8].upper()}"
            risk_level = 3
            bid = models.Bid(
                bid_code=bid_code, contract_title=opp.title, agency_name=opp.agency_name,
                agency_type=opp.agency_type or "federal",
                contract_type=opp.contract_type or "service", delivery_distance_miles=0,
                deadline_date=opp.due_date.date() if opp.due_date else date.today(),
                urgency_level=3, competition_level="medium", risk_level=risk_level,
                desired_profit_mode="balanced", status="draft",
                notes=f"Auto-generated by Bid Autopilot from {opp.opp_code}",
            )
            db.add(bid)
            db.flush()
            opp.converted_bid_id = bid.id
            opp.status = "converted"
            db.commit()
            db.refresh(bid)
            steps_completed.append("bid")
        except Exception as e:
            errors.append(f"Bid: {str(e)}")

    # Generate Proposal
    proposal_path = None
    try:
        analysis_context = ""
        if opp.ai_summary:
            analysis_context += f"\nSummary: {opp.ai_summary[:800]}"
        if opp.ai_bid_strategy:
            analysis_context += f"\nStrategy: {opp.ai_bid_strategy[:500]}"

        pricing_context = ""
        proposal_text = None

        try:
            from ..claude_ai import call_claude
            proposal_text = call_claude(
                prompt=f"Write a complete government proposal for:\n{context[:4000]}\n\nAnalysis:\n{analysis_context}\n\nCompany: {profile_text[:2000]}\n\nSections: Cover Letter, Executive Summary, Technical Approach, Management Plan, Past Performance, Staffing, QA, Pricing Narrative. Use \"{company_name}\" throughout.",
                system_instruction=f"You are an expert government proposal writer for {company_name}.",
                max_tokens=8000,
            )
        except Exception:
            pass

        if not proposal_text:
            proposal_text = _call_openai(
                f"You are an expert proposal writer for {company_name}.",
                f"Write a complete government proposal for:\n{context[:3000]}\n\nProfile:\n{profile_text[:2000]}",
                json_mode=False,
            )

        if proposal_text:
            profile_obj = u.profile if u else None
            from ..proposal_generator import generate_pdf_proposal, generate_docx_proposal
            if format_type == "docx":
                proposal_path = generate_docx_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid)
            else:
                proposal_path = generate_pdf_proposal(opp, proposal_text, company_name, profile=profile_obj, bid=bid)
            steps_completed.append("proposal")

            if not opp.capture_notes:
                opp.capture_notes = ""
            opp.capture_notes = f"autopilot_proposal_path={proposal_path}\n" + (opp.capture_notes or "")
            db.commit()
    except Exception as e:
        errors.append(f"Proposal: {str(e)}")

    return {
        "ok": True,
        "opp_id": opp.id,
        "opp_code": opp.opp_code,
        "bid_id": bid.id if bid else None,
        "bid_code": bid.bid_code if bid else None,
        "steps_completed": steps_completed,
        "steps_total": 7,
        "errors": errors,
        "analysis": {
            "recommendation": opp.ai_bid_recommendation,
            "confidence": opp.ai_confidence_score,
        },
        "war_room": {
            "win_probability": war_room.get("win_probability", {}).get("score", 0) if war_room else 0,
            "bottom_line": war_room.get("bottom_line", "") if war_room else "",
        },
        "compliance": {
            "score": compliance_matrix.get("compliance_score", 0) if compliance_matrix else 0,
        },
        "has_proposal": proposal_path is not None,
        "proposal_download": f"/opportunities/{opp.id}/autopilot-download?format={format_type}" if proposal_path else None,
    }


@router.get("/{opp_id}/autopilot-download")
def autopilot_download(
    opp_id: int,
    format: str = "pdf",
    user: str = Depends(require_auth),
    db: Session = Depends(get_db),
):
    """Download the proposal generated by Bid Autopilot."""
    opp = db.query(models.Opportunity).filter(models.Opportunity.id == opp_id).first()
    if not opp:
        raise HTTPException(404, "Opportunity not found")

    # Find proposal path from capture_notes
    proposal_path = None
    if opp.capture_notes:
        for line in opp.capture_notes.split("\n"):
            if line.startswith("autopilot_proposal_path="):
                proposal_path = line.split("=", 1)[1].strip()
                break

    if not proposal_path or not os.path.exists(proposal_path):
        raise HTTPException(404, "Proposal file not found. Run Autopilot again.")

    media = "application/pdf" if format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(
        proposal_path,
        media_type=media,
        filename=f"Proposal-{opp.opp_code}.{format}",
        headers={"Content-Disposition": f"attachment; filename=Proposal-{opp.opp_code}.{format}"},
    )

