# app/services/exporter.py
from __future__ import annotations

import csv
import io
import json
from typing import Dict, Any, List, Iterable, Optional

from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


# ----------------------------
# Helpers
# ----------------------------

def _safe_str(x: Any) -> str:
    if x is None:
        return ""
    return str(x)

def _safe_float(x: Any) -> Optional[float]:
    try:
        if x is None:
            return None
        return float(x)
    except Exception:
        return None

def _money(x: Any) -> str:
    v = _safe_float(x)
    if v is None:
        return ""
    return f"{v:,.2f}"

def _jsonish(x: Any) -> str:
    if isinstance(x, (dict, list)):
        return json.dumps(x, ensure_ascii=False)
    return _safe_str(x)

def _wrap(text: str, width: int) -> List[str]:
    words = (text or "").split()
    lines: List[str] = []
    cur = ""
    for w in words:
        if not cur:
            cur = w
            continue
        if len(cur) + 1 + len(w) <= width:
            cur += " " + w
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines

def _as_attachment(filename: str) -> Dict[str, str]:
    return {"Content-Disposition": f'attachment; filename="{filename}"'}


# ----------------------------
# CSV Export
# ----------------------------

def export_csv_summary(payload: Dict[str, Any]) -> StreamingResponse:
    """
    Exports a flat key/value CSV summary of the frozen version payload.
    """
    buf = io.StringIO(newline="")
    w = csv.writer(buf)
    w.writerow(["field", "value"])
    for k, v in payload.items():
        w.writerow([k, _jsonish(v)])

    data = buf.getvalue().encode("utf-8-sig")  # Excel-friendly UTF-8 BOM
    return StreamingResponse(
        io.BytesIO(data),
        media_type="text/csv; charset=utf-8",
        headers=_as_attachment("summary.csv"),
    )


# ----------------------------
# DOCX Export
# ----------------------------

def export_docx_from_payload(payload: Dict[str, Any]) -> StreamingResponse:
    """
    Exports a detailed DOCX bid package from the frozen BidVersion payload.
    Includes pricing summary + justification + tables (items/labor/equipment).
    """
    doc = Document()

    doc.add_heading("SentriBiD — Detailed Bid Package", 0)
    doc.add_paragraph(f"Bid: {payload.get('bid_code','')}  v{payload.get('version_no','')}")
    doc.add_paragraph(f"Contract: {payload.get('contract_title','')}")
    doc.add_paragraph(f"Agency: {payload.get('agency_name','')} ({payload.get('agency_type','')})")
    doc.add_paragraph("")

    # Pricing Summary
    doc.add_heading("Pricing Summary", level=1)
    summary_fields = [
        ("Selected Mode", payload.get("selected_mode")),
        ("Final Bid Price", _money(payload.get("final_bid_price"))),
        ("Profit Amount", _money(payload.get("profit_amount"))),
        ("Profit %", _safe_str(payload.get("profit_pct"))),
        ("Win Score", _safe_str(payload.get("win_score"))),
        ("True Cost", _money(payload.get("true_cost"))),
        ("Risk Buffer", _money(payload.get("risk_buffer"))),
        ("Adjusted Cost", _money(payload.get("adjusted_cost"))),
        ("Margin %", _safe_str(payload.get("margin_pct"))),
    ]
    for k, v in summary_fields:
        doc.add_paragraph(f"{k}: {v if v is not None else ''}")

    # Justification
    doc.add_heading("Justification", level=1)
    doc.add_paragraph(_safe_str(payload.get("justification_text", "")))

    # Items table
    items = payload.get("items") or []
    if items:
        doc.add_heading("Items", level=1)
        table = doc.add_table(rows=1, cols=7)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Description"
        hdr[2].text = "Qty"
        hdr[3].text = "Unit Cost"
        hdr[4].text = "Line Total"
        hdr[5].text = "Supplier"
        hdr[6].text = "Risk"

        for it in items:
            row = table.add_row().cells
            row[0].text = _safe_str(it.get("name"))
            row[1].text = _safe_str(it.get("description"))
            row[2].text = _safe_str(it.get("quantity"))
            row[3].text = _money(it.get("unit_cost"))
            row[4].text = _money(it.get("line_total"))
            row[5].text = _safe_str(it.get("supplier_name"))
            row[6].text = "Yes" if bool(it.get("risk_flag")) else "No"

    # Labor table
    labor = payload.get("labor") or []
    if labor:
        doc.add_heading("Labor", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Type"
        hdr[1].text = "Hourly Rate"
        hdr[2].text = "Hours"
        hdr[3].text = "Workers"
        hdr[4].text = "Line Total"
        hdr[5].text = "Notes"

        for ln in labor:
            row = table.add_row().cells
            row[0].text = _safe_str(ln.get("labor_type"))
            row[1].text = _money(ln.get("hourly_rate"))
            row[2].text = _safe_str(ln.get("hours"))
            row[3].text = _safe_str(ln.get("workers"))
            row[4].text = _money(ln.get("line_total"))
            row[5].text = ""

    # Equipment table
    equipment = payload.get("equipment") or []
    if equipment:
        doc.add_heading("Equipment", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Equipment"
        hdr[1].text = "Rental Cost"
        hdr[2].text = "Days"
        hdr[3].text = "Operator Required"
        hdr[4].text = "Operator Cost"
        hdr[5].text = "Total"

        for eq in equipment:
            rental_cost = _safe_float(eq.get("rental_cost")) or 0.0
            days = int(_safe_float(eq.get("rental_days")) or 1)
            op_cost = (_safe_float(eq.get("operator_cost")) or 0.0) if bool(eq.get("operator_required")) else 0.0
            total = (rental_cost * max(days, 1)) + op_cost

            row = table.add_row().cells
            row[0].text = _safe_str(eq.get("equipment_name"))
            row[1].text = _money(rental_cost)
            row[2].text = _safe_str(days)
            row[3].text = "Yes" if bool(eq.get("operator_required")) else "No"
            row[4].text = _money(op_cost)
            row[5].text = _money(total)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=_as_attachment("bid_package.docx"),
    )


# ----------------------------
# PDF Export
# ----------------------------

def export_pdf_from_payload(payload: Dict[str, Any]) -> StreamingResponse:
    """
    Exports a readable PDF bid package:
    - Header + pricing summary
    - Justification (wrapped)
    - Items/Labor/Equipment sections
    """
    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=letter)
    width, height = letter
    left = 72
    y = height - 72

    def new_page():
        nonlocal y
        c.showPage()
        y = height - 72

    def line(txt: str = "", dy: int = 14, bold: bool = False, size: int = 11):
        nonlocal y
        if y < 72:
            new_page()
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(left, y, _safe_str(txt)[:160])
        y -= dy

    def spacer(dy: int = 10):
        nonlocal y
        y -= dy
        if y < 72:
            new_page()

    # Header
    line("SentriBiD — Detailed Bid Package", bold=True, dy=18, size=14)
    line(f"Bid: {payload.get('bid_code','')}  v{payload.get('version_no','')}", dy=16)
    line(f"Contract: {payload.get('contract_title','')}")
    line(f"Agency: {payload.get('agency_name','')} ({payload.get('agency_type','')})")
    spacer()

    # Pricing Summary
    line("Pricing Summary", bold=True, dy=16, size=12)
    summary_lines = [
        ("Selected Mode", payload.get("selected_mode")),
        ("Final Bid Price", _money(payload.get("final_bid_price"))),
        ("Profit Amount", _money(payload.get("profit_amount"))),
        ("Profit %", payload.get("profit_pct")),
        ("Win Score", payload.get("win_score")),
        ("True Cost", _money(payload.get("true_cost"))),
        ("Risk Buffer", _money(payload.get("risk_buffer"))),
        ("Adjusted Cost", _money(payload.get("adjusted_cost"))),
        ("Margin %", payload.get("margin_pct")),
    ]
    for k, v in summary_lines:
        line(f"{k}: {_safe_str(v)}")
    spacer()

    # Justification
    line("Justification", bold=True, dy=16, size=12)
    for chunk in _wrap(_safe_str(payload.get("justification_text", "")), 95):
        line(chunk)
    spacer()

    # Items
    items = payload.get("items") or []
    if items:
        line("Items", bold=True, dy=16, size=12)
        line("Name | Qty | Unit Cost | Line Total | Supplier | Risk", dy=14, bold=True)
        for it in items:
            txt = (
                f"{_safe_str(it.get('name'))[:30]} | "
                f"{_safe_str(it.get('quantity'))} | "
                f"{_money(it.get('unit_cost'))} | "
                f"{_money(it.get('line_total'))} | "
                f"{_safe_str(it.get('supplier_name'))[:18]} | "
                f"{'Y' if bool(it.get('risk_flag')) else 'N'}"
            )
            line(txt)
        spacer()

    # Labor
    labor = payload.get("labor") or []
    if labor:
        line("Labor", bold=True, dy=16, size=12)
        line("Type | Rate | Hours | Workers | Line Total", dy=14, bold=True)
        for ln in labor:
            txt = (
                f"{_safe_str(ln.get('labor_type'))[:18]} | "
                f"{_money(ln.get('hourly_rate'))} | "
                f"{_safe_str(ln.get('hours'))} | "
                f"{_safe_str(ln.get('workers'))} | "
                f"{_money(ln.get('line_total'))}"
            )
            line(txt)
        spacer()

    # Equipment
    equipment = payload.get("equipment") or []
    if equipment:
        line("Equipment", bold=True, dy=16, size=12)
        line("Equipment | Rental | Days | Operator | Operator Cost | Total", dy=14, bold=True)
        for eq in equipment:
            rental_cost = _safe_float(eq.get("rental_cost")) or 0.0
            days = int(_safe_float(eq.get("rental_days")) or 1)
            op_required = bool(eq.get("operator_required"))
            op_cost = (_safe_float(eq.get("operator_cost")) or 0.0) if op_required else 0.0
            total = (rental_cost * max(days, 1)) + op_cost

            txt = (
                f"{_safe_str(eq.get('equipment_name'))[:22]} | "
                f"{_money(rental_cost)} | "
                f"{days} | "
                f"{'Y' if op_required else 'N'} | "
                f"{_money(op_cost)} | "
                f"{_money(total)}"
            )
            line(txt)
        spacer()

    c.save()
    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/pdf",
        headers=_as_attachment("bid_package.pdf"),
    )
